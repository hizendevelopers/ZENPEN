from __future__ import annotations

import base64
import hashlib
import importlib
import importlib.util
import json
import logging
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
from collections import Counter
from concurrent.futures import ThreadPoolExecutor
from contextlib import contextmanager
from datetime import datetime, timezone
from functools import lru_cache
from html import unescape
from io import BytesIO
from pathlib import Path
from typing import Optional
from urllib.parse import urlparse

import nltk
import numpy as np
import httpx
from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, StreamingResponse
from pydantic import BaseModel
from starlette.staticfiles import StaticFiles
from nltk.tokenize import sent_tokenize
from dotenv import load_dotenv
from rq.job import Job

from backend.queueing import fetch_job, queue_is_available, queue_is_configured, workers_available
from backend.url_jobs import enqueue_url_analysis

BASE_DIR = Path(__file__).resolve().parent.parent
FRONTEND_DIR = BASE_DIR / "frontend"
HISTORY_FILE = BASE_DIR / "backend" / "history.json"
WHISPER_CACHE_DIR = BASE_DIR / "backend" / ".cache" / "whisper"
TRANSCRIPT_CACHE_DIR = BASE_DIR / "backend" / ".cache" / "transcripts"
SOURCE_CACHE_DIR = BASE_DIR / "backend" / ".cache" / "sources"
SUMMARY_CACHE_DIR = BASE_DIR / "backend" / ".cache" / "summaries"
ANALYSIS_CACHE_DIR = BASE_DIR / "backend" / ".cache" / "analysis"
ARTICLE_CACHE_DIR = BASE_DIR / "backend" / ".cache" / "articles"
ENV_FILE = BASE_DIR / ".env"

os.environ.setdefault("HF_HUB_DISABLE_PROGRESS_BARS", "1")
os.environ.setdefault("TOKENIZERS_PARALLELISM", "false")
load_dotenv(ENV_FILE)

app = FastAPI(title="Media Transcriber & Analyzer")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
if FRONTEND_DIR.exists():
    app.mount("/static", StaticFiles(directory=str(FRONTEND_DIR)), name="static")

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
SUPABASE_URL = os.getenv("SUPABASE_URL", "").rstrip("/")
SUPABASE_PUBLISHABLE_KEY = os.getenv("SUPABASE_PUBLISHABLE_KEY", "")
SUPABASE_SECRET_KEY = os.getenv("SUPABASE_SECRET_KEY", "")
YOUTUBE_COOKIES_B64 = os.getenv("YOUTUBE_COOKIES_B64", "").strip()
APIFY_TOKEN = os.getenv("APIFY_TOKEN", "").strip()
REDIS_URL = os.getenv("REDIS_URL", "").strip()
YOUTUBE_PROXY_URL = os.getenv("YOUTUBE_PROXY_URL", "").strip()
YOUTUBE_PROXY_HTTP = os.getenv("YOUTUBE_PROXY_HTTP", "").strip()
YOUTUBE_PROXY_HTTPS = os.getenv("YOUTUBE_PROXY_HTTPS", "").strip()
FAST_ANALYSIS_MODE = os.getenv("FAST_ANALYSIS_MODE", "true").strip().lower() not in {"0", "false", "no", "off"}
FAST_ANALYSIS_TRANSCRIPT_LIMIT = int(os.getenv("FAST_ANALYSIS_TRANSCRIPT_LIMIT", "5000"))
FAST_ANALYSIS_CHUNK_LIMIT = int(os.getenv("FAST_ANALYSIS_CHUNK_LIMIT", "6"))
ENABLE_BACKGROUND_URL_JOBS = os.getenv("ENABLE_BACKGROUND_URL_JOBS", "true").strip().lower() in {"1", "true", "yes", "on"}
ENABLE_REMOTE_MEDIA_DOWNLOAD_FALLBACK = os.getenv("ENABLE_REMOTE_MEDIA_DOWNLOAD_FALLBACK", "false").strip().lower() in {"1", "true", "yes", "on"}
TRANSCRIPTION_CHUNK_SECONDS = int(os.getenv("TRANSCRIPTION_CHUNK_SECONDS", "600"))
LONG_TRANSCRIPT_THRESHOLD = int(os.getenv("LONG_TRANSCRIPT_THRESHOLD", "7000"))
SUMMARY_CONCURRENCY = max(1, int(os.getenv("SUMMARY_CONCURRENCY", "3")))
ENABLE_DEEP_ARTICLE_REFINEMENT = os.getenv("ENABLE_DEEP_ARTICLE_REFINEMENT", "false").strip().lower() in {"1", "true", "yes", "on"}
SOURCE_CACHE_TTL_SECONDS = int(os.getenv("SOURCE_CACHE_TTL_SECONDS", str(60 * 60 * 12)))
YOUTUBE_DIRECT_ANALYSIS_TIMEOUT = int(os.getenv("YOUTUBE_DIRECT_ANALYSIS_TIMEOUT", "120"))
YOUTUBE_TRANSCRIPT_TIMEOUT = int(os.getenv("YOUTUBE_TRANSCRIPT_TIMEOUT", "30"))
YOUTUBE_SUBTITLE_TIMEOUT = int(os.getenv("YOUTUBE_SUBTITLE_TIMEOUT", "30"))
YOUTUBE_FALLBACK_TIMEOUT = int(os.getenv("YOUTUBE_FALLBACK_TIMEOUT", "180"))
ANALYSIS_JOB_MAX_SECONDS = int(os.getenv("ANALYSIS_JOB_MAX_SECONDS", "480"))
ANALYSIS_JOB_STALE_SECONDS = int(os.getenv("ANALYSIS_JOB_STALE_SECONDS", "150"))
GEMINI_BACKOFF_UNTIL = 0.0

logger = logging.getLogger("zenpen.performance")
if not logger.handlers:
    handler = logging.StreamHandler()
    handler.setFormatter(logging.Formatter("%(asctime)s %(levelname)s [%(name)s] %(message)s"))
    logger.addHandler(handler)
logger.setLevel(logging.INFO)

WHISPER_MODEL = None
EMBEDDER = None
GENAI_CLIENT = None
FAISS_MODULE = None
YT_DLP_MODULE = None
VIDEO_FILE_CLIP_CLASS = None
SENTENCE_TRANSFORMER_CLASS = None
WHISPER_MODULE = None
GENAI_MODULE = None
APIFY_CLIENT_CLASS = None
YOUTUBE_TRANSCRIPT_API_CLASS = None


class ArticlesRequest(BaseModel):
    headline: str
    summary: str
    topics: list[str] = []
    selected_topics: list[str]
    selected_topic_details: dict[str, object] | None = None
    article_count: int = 1
    article_type: str = "Blog Article"
    target_audience: str = "General readers"
    source_context: str = ""
    source_cache_key: str = ""


class PublishArticleRequest(BaseModel):
    headline: str
    summary: str
    topics: list[str] = []
    selected_topics: list[str] = []
    articles: list[dict[str, object]]
    source_type: str = "url"
    source_url: Optional[str] = None
    source_file_name: Optional[str] = None
    source_mime_type: Optional[str] = None
    query: str = "Give breaking news and main points"


class ExportArticleRequest(BaseModel):
    title: str
    topic: str
    content_html: str
    format: str = "txt"


class AnalyzeYouTubeRequest(BaseModel):
    url: str
    query: str = "Give breaking news and main points"


class StageProfiler:
    def __init__(self, label: str):
        self.label = label
        self.started_at = time.perf_counter()
        self.steps: list[tuple[str, float]] = []

    @contextmanager
    def stage(self, name: str):
        step_start = time.perf_counter()
        try:
            yield
        finally:
            duration = time.perf_counter() - step_start
            self.steps.append((name, duration))
            logger.info("%s | %s | %.2fs", self.label, name, duration)

    def log_total(self) -> None:
        total = time.perf_counter() - self.started_at
        breakdown = ", ".join(f"{name}={duration:.2f}s" for name, duration in self.steps) or "no steps"
        logger.info("%s | total=%.2fs | %s", self.label, total, breakdown)


class QuietYtdlpLogger:
    def debug(self, msg: str) -> None:
        return None

    def warning(self, msg: str) -> None:
        return None

    def error(self, msg: str) -> None:
        return None


class SignupRequest(BaseModel):
    name: str
    email: str
    password: str


class LoginRequest(BaseModel):
    email: str
    password: str


COMMON_TOPIC_STOPWORDS = {
    "about", "after", "again", "also", "amid", "among", "around", "being", "below", "between",
    "both", "could", "first", "from", "further", "here", "into", "just", "main", "more", "most",
    "news", "only", "other", "over", "same", "some", "such", "than", "that", "their", "there",
    "these", "they", "this", "those", "through", "under", "using", "very", "what", "when", "where",
    "which", "while", "with", "would", "your", "have", "were", "been", "them", "because", "content",
    "video", "transcript", "summary", "give", "breaking", "points", "said", "says", "show", "kind",
    "then", "taken", "single", "generation", "musical", "story", "system", "shown", "last", "this", "that",
}

ALLOWED_ARTICLE_TAGS = {"h2", "h3", "p", "ul", "li", "strong", "br"}
BANNED_ARTICLE_HEADINGS = {
    "why this topic stands out",
    "key developments",
    "what it suggests",
    "selected topic",
    "source material",
}
GENERIC_TOPIC_TITLES = {
    "topic", "main topic", "summary", "overview", "news", "update", "speaker", "lyrics",
    "strong", "explicit", "article", "content", "message", "video", "story", "analysis",
    "commitment",
}
SUSPICIOUS_ANALYSIS_TITLES = {
    "https", "http", "instagram", "youtube", "mooroo", "channel", "credits", "link", "links",
}
SUPPORTED_UPLOAD_MIME_PREFIXES = ("video/", "audio/")
SUPPORTED_MEDIA_EXTENSIONS = {".mp3", ".wav", ".m4a", ".aac", ".ogg", ".flac", ".mp4", ".mov", ".mkv", ".avi", ".webm"}
UNSUPPORTED_DIRECT_SOURCE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".svg", ".bmp", ".ico"}
ARTICLE_TYPE_INSTRUCTIONS = {
    "Blog Article": "Write a polished editorial blog article with a clear narrative arc, informative tone, and reader-friendly explanations.",
    "News Article": "Write in a neutral newsroom style with a sharp headline, an efficient lead paragraph, key details first, then context and developments.",
    "How-to Guide": "Write as a practical step-by-step guide with clear instructional subheadings, helpful sequencing, and concise examples.",
    "Listicle": "Write as a high-quality list-based article with dynamic numbered sections and crisp transitions, not as a shallow clickbait list.",
    "Review": "Write as a balanced review that weighs strengths, weaknesses, evidence, and audience relevance without sounding promotional.",
    "SEO Article": "Write as a search-optimized long-form article with clean structure, semantic keyword coverage, strong intent match, and direct answers.",
}
FILLER_PATTERNS = [
    r"\b(?:um+|uh+|erm|hmm+)\b",
    r"\b(?:you know|i mean|kind of|sort of)\b",
    r"\[(?:music|applause|laughter|noise)[^\]]*\]",
    r"\(?(?:music|applause|laughter|noise)\)?",
]
RAW_ERROR_PATTERNS = [
    r"\b\d{3}\b",
    r"traceback",
    r"runtimeerror",
    r"googleapierror",
    r"httpx\.",
    r"youtube transcript fallback failed",
]


def dependency_status() -> dict[str, bool | str]:
    return {
        "python": sys.executable,
        "yt_dlp": module_available("yt_dlp"),
        "whisper": module_available("whisper"),
        "moviepy": module_available("moviepy"),
        "faiss": module_available("faiss"),
        "sentence_transformers": module_available("sentence_transformers"),
        "gemini": module_available("google.genai"),
        "ffmpeg": shutil.which("ffmpeg") is not None,
        "gemini_api_key_configured": bool(GEMINI_API_KEY),
        "apify_configured": bool(APIFY_TOKEN),
        "youtube_transcript_api": module_available("youtube_transcript_api"),
        "supabase_configured": supabase_is_configured(),
        "redis_queue_configured": queue_is_configured(),
        "redis_queue_available": queue_is_available(),
        "background_url_jobs_enabled": background_url_jobs_available(),
        "remote_media_download_fallback_enabled": ENABLE_REMOTE_MEDIA_DOWNLOAD_FALLBACK,
        "remote_media_fallback_available": remote_media_fallback_available(),
        "youtube_proxy_configured": bool(get_proxy_url("http") or get_proxy_url("https")),
    }


def module_available(module_name: str) -> bool:
    try:
        return importlib.util.find_spec(module_name) is not None
    except Exception:
        return False


def get_genai_module():
    global GENAI_MODULE
    if GENAI_MODULE is None:
        try:
            GENAI_MODULE = importlib.import_module("google.genai")
        except Exception as exc:
            raise RuntimeError(build_missing_dependency_message("google-genai")) from exc
    return GENAI_MODULE


def get_whisper_module():
    global WHISPER_MODULE
    if WHISPER_MODULE is None:
        try:
            WHISPER_MODULE = importlib.import_module("whisper")
        except Exception as exc:
            raise RuntimeError(build_missing_dependency_message("Whisper")) from exc
    return WHISPER_MODULE


def get_sentence_transformer_class():
    global SENTENCE_TRANSFORMER_CLASS
    if SENTENCE_TRANSFORMER_CLASS is None:
        try:
            module = importlib.import_module("sentence_transformers")
            SENTENCE_TRANSFORMER_CLASS = getattr(module, "SentenceTransformer")
        except Exception as exc:
            raise RuntimeError(build_missing_dependency_message("sentence-transformers")) from exc
    return SENTENCE_TRANSFORMER_CLASS


def get_yt_dlp_module():
    global YT_DLP_MODULE
    if YT_DLP_MODULE is None:
        try:
            YT_DLP_MODULE = importlib.import_module("yt_dlp")
        except Exception as exc:
            raise RuntimeError(build_missing_dependency_message("yt-dlp")) from exc
    return YT_DLP_MODULE


def get_video_file_clip_class():
    global VIDEO_FILE_CLIP_CLASS
    if VIDEO_FILE_CLIP_CLASS is None:
        try:
            module = importlib.import_module("moviepy.editor")
            VIDEO_FILE_CLIP_CLASS = getattr(module, "VideoFileClip")
        except Exception:
            try:
                module = importlib.import_module("moviepy")
                VIDEO_FILE_CLIP_CLASS = getattr(module, "VideoFileClip")
            except Exception as exc:
                raise RuntimeError(build_missing_dependency_message("moviepy")) from exc
    return VIDEO_FILE_CLIP_CLASS


def get_faiss_module():
    global FAISS_MODULE
    if FAISS_MODULE is None:
        try:
            FAISS_MODULE = importlib.import_module("faiss")
        except Exception as exc:
            raise RuntimeError(build_missing_dependency_message("faiss-cpu")) from exc
    return FAISS_MODULE


def get_apify_client_class():
    global APIFY_CLIENT_CLASS
    if APIFY_CLIENT_CLASS is None:
        try:
            module = importlib.import_module("apify_client")
            APIFY_CLIENT_CLASS = getattr(module, "ApifyClient")
        except Exception as exc:
            raise RuntimeError(build_missing_dependency_message("apify-client")) from exc
    return APIFY_CLIENT_CLASS


def get_youtube_transcript_api_class():
    global YOUTUBE_TRANSCRIPT_API_CLASS
    if YOUTUBE_TRANSCRIPT_API_CLASS is None:
        try:
            module = importlib.import_module("youtube_transcript_api")
            YOUTUBE_TRANSCRIPT_API_CLASS = getattr(module, "YouTubeTranscriptApi")
        except Exception as exc:
            raise RuntimeError(build_missing_dependency_message("youtube-transcript-api")) from exc
    return YOUTUBE_TRANSCRIPT_API_CLASS


def source_kind_from_url(url: str) -> str:
    parsed = urlparse(url)
    path = (parsed.path or "").lower()
    suffix = Path(path).suffix.lower()
    if is_youtube_url(url):
        return "youtube"
    if suffix in UNSUPPORTED_DIRECT_SOURCE_EXTENSIONS:
        return "unsupported"
    if suffix in SUPPORTED_MEDIA_EXTENSIONS:
        return "media-url"
    return "web-url"


def map_public_error_message(raw_message: str, *, context: str = "analysis") -> str:
    message = (raw_message or "").strip()
    lower = message.lower()
    if not message:
        return "Something went wrong while processing your request. Please try again."
    if "password" in lower and "match" in lower:
        return "Passwords do not match. Please check them and try again."
    if "invalid email or password" in lower:
        return "We could not log you in with those credentials."
    if "could not create your account" in lower:
        return "We could not create your account right now. Please try again."
    if "please provide a url or upload" in lower:
        return "Please add a valid source before continuing."
    if "unsupported" in lower and "image" in lower:
        return "Image links are not supported yet. Please use a webpage, YouTube/video URL, or upload an audio/video file."
    if "unsupported" in lower and "file" in lower:
        return "That file type is not supported. Please upload an audio or video file."
    if "failed to load audio" in lower or ("ffmpeg" in lower and "invalid data found when processing input" in lower):
        return "That file could not be processed as audio or video. Please upload a supported media file."
    if "video could not be downloaded" in lower:
        return "Video could not be downloaded. The source may be private, unavailable, or blocking automated access."
    if "gemini could not analyze the video directly" in lower:
        return "Gemini could not analyze the video directly."
    if "transcription failed" in lower:
        return "Transcription failed. Please try another source or upload a clearer audio/video file."
    if "website blocked direct content extraction" in lower:
        return "The website blocked automated access. Please try another public URL."
    if "youtube" in lower and "requires authenticated cookies" in lower:
        return "This YouTube video cannot be accessed from the server right now. Please try another video or upload the file directly."
    if "rate" in lower and "limit" in lower:
        return "The AI service is busy right now. Please wait a moment and try again."
    if "timeout" in lower:
        if "youtube metadata analysis" in lower or "youtube transcript fetch" in lower or "youtube subtitle fetch" in lower:
            return "Direct video analysis timed out. Please retry or try another video."
        if "transcription fallback" in lower or "video download" in lower:
            return "Fallback processing timed out before the video could be transcribed."
        return "The request took too long to complete. Please try again."
    if "private" in lower or "unavailable" in lower:
        return "Source is private or unavailable."
    if "unsupported video format" in lower:
        return "Please upload a supported audio or video file."
    if "gemini" in lower or "googleapierror" in lower:
        return "The AI writing service could not complete this request right now. Please try again."
    if any(re.search(pattern, lower) for pattern in RAW_ERROR_PATTERNS):
        if context == "login":
            return "We could not complete the login request right now."
        return "We could not complete the request right now. Please try again."
    return message if len(message) < 220 else "We could not complete the request right now. Please try again."


def get_beautiful_soup_class():
    try:
        module = importlib.import_module("bs4")
        return getattr(module, "BeautifulSoup")
    except Exception as exc:
        raise RuntimeError(build_missing_dependency_message("beautifulsoup4")) from exc


def supabase_is_configured() -> bool:
    return bool(SUPABASE_URL and SUPABASE_PUBLISHABLE_KEY and SUPABASE_SECRET_KEY)


def get_api_config() -> dict[str, object]:
    return {
        "supabase": {
            "enabled": bool(SUPABASE_URL and SUPABASE_PUBLISHABLE_KEY),
            "url": SUPABASE_URL or None,
            "publishableKey": SUPABASE_PUBLISHABLE_KEY or None,
        },
        "queue": {
            "enabled": background_url_jobs_available(),
        },
        "supportedSources": ["web-url", "youtube", "media-url", "uploaded-audio", "uploaded-video"],
    }


def background_url_jobs_available() -> bool:
    return ENABLE_BACKGROUND_URL_JOBS and queue_is_available() and workers_available()


def get_bearer_token(request: Request) -> Optional[str]:
    authorization = request.headers.get("Authorization", "").strip()
    if not authorization.lower().startswith("bearer "):
        return None
    token = authorization.split(" ", 1)[1].strip()
    return token or None


def supabase_auth_headers(token: str) -> dict[str, str]:
    return {
        "apikey": SUPABASE_PUBLISHABLE_KEY,
        "Authorization": f"Bearer {token}",
    }


def supabase_service_headers() -> dict[str, str]:
    return {
        "apikey": SUPABASE_SECRET_KEY,
        "Authorization": f"Bearer {SUPABASE_SECRET_KEY}",
        "Content-Type": "application/json",
        "Prefer": "return=representation",
    }


def supabase_public_headers() -> dict[str, str]:
    return {
        "apikey": SUPABASE_PUBLISHABLE_KEY,
        "Content-Type": "application/json",
    }


def supabase_auth_request(
    method: str,
    path: str,
    *,
    params: Optional[dict[str, str]] = None,
    json_body: Optional[object] = None,
    admin: bool = False,
) -> httpx.Response:
    if not supabase_is_configured():
        raise RuntimeError("Supabase is not configured")

    headers = supabase_service_headers() if admin else supabase_public_headers()
    response = httpx.request(
        method,
        f"{SUPABASE_URL}/auth/v1/{path}",
        headers=headers,
        params=params,
        json=json_body,
        timeout=20.0,
    )
    return response


def extract_supabase_error(response: httpx.Response, default_message: str) -> str:
    try:
        payload = response.json()
    except Exception:
        return default_message

    if isinstance(payload, dict):
        for key in ("msg", "message", "error_description", "error"):
            value = payload.get(key)
            if value:
                return str(value)
    return default_message


def create_supabase_user(name: str, email: str, password: str) -> dict[str, object]:
    response = supabase_auth_request(
        "POST",
        "admin/users",
        json_body={
            "email": email,
            "password": password,
            "email_confirm": True,
            "user_metadata": {
                "full_name": name,
            },
        },
        admin=True,
    )

    if response.status_code >= 400:
        raise HTTPException(
            status_code=400,
            detail=extract_supabase_error(response, "Could not create your account."),
        )

    payload = response.json()
    return payload if isinstance(payload, dict) else {}


def sign_in_supabase_user(email: str, password: str) -> dict[str, object]:
    response = supabase_auth_request(
        "POST",
        "token",
        params={"grant_type": "password"},
        json_body={
            "email": email,
            "password": password,
        },
    )

    if response.status_code >= 400:
        raise HTTPException(
            status_code=400,
            detail=extract_supabase_error(response, "Invalid email or password."),
        )

    payload = response.json()
    return payload if isinstance(payload, dict) else {}


def resolve_supabase_user(request: Request) -> Optional[dict[str, object]]:
    if not supabase_is_configured():
        return None

    token = get_bearer_token(request)
    if not token:
        return None

    try:
        response = httpx.get(
            f"{SUPABASE_URL}/auth/v1/user",
            headers=supabase_auth_headers(token),
            timeout=15.0,
        )
        if response.status_code != 200:
            return None
        payload = response.json()
        return payload if isinstance(payload, dict) and payload.get("id") else None
    except Exception:
        return None


def supabase_rest_request(
    method: str,
    path: str,
    *,
    params: Optional[dict[str, str]] = None,
    json_body: Optional[object] = None,
) -> object:
    if not supabase_is_configured():
        raise RuntimeError("Supabase is not configured")

    response = httpx.request(
        method,
        f"{SUPABASE_URL}/rest/v1/{path}",
        headers=supabase_service_headers(),
        params=params,
        json=json_body,
        timeout=20.0,
    )
    response.raise_for_status()
    if not response.content:
        return None
    return response.json()


@lru_cache(maxsize=8)
def get_product_id_by_slug(slug: str) -> Optional[str]:
    rows = supabase_rest_request(
        "GET",
        "products",
        params={
            "slug": f"eq.{slug}",
            "select": "id",
            "limit": "1",
        },
    )
    if isinstance(rows, list) and rows:
        product_id = rows[0].get("id")
        return str(product_id) if product_id else None
    return None


def load_history_from_supabase(user_id: str) -> list[dict[str, object]]:
    rows = supabase_rest_request(
        "GET",
        "analysis_jobs",
        params={
            "user_id": f"eq.{user_id}",
            "select": "id,source_type,source_url,source_file_name,headline,summary,created_at",
            "order": "created_at.desc",
            "limit": "10",
        },
    )
    history_items: list[dict[str, object]] = []
    if not isinstance(rows, list):
        return history_items

    for row in rows:
        source_type = row.get("source_type") or "url"
        source_label = "youtube-url" if source_type == "url" else "uploaded-file"
        history_items.append(
            {
                "id": row.get("id"),
                "source": source_label,
                "headline": row.get("headline") or "Media summary",
                "summary": row.get("summary") or "",
                "created_at": row.get("created_at"),
            }
        )
    return history_items


def persist_analysis_to_supabase(
    *,
    user_id: str,
    source_type: str,
    source_url: Optional[str],
    source_file_name: Optional[str],
    source_mime_type: Optional[str],
    query: str,
    result: dict[str, object],
    selected_topics: list[str],
) -> None:
    product_id = get_product_id_by_slug("article-generator")
    if not product_id:
        raise RuntimeError("The article-generator product row is missing in Supabase.")

    headline = str(result.get("headline") or "Media summary")
    summary = str(result.get("summary") or "")
    topics = result.get("topics") if isinstance(result.get("topics"), list) else []
    articles = result.get("articles") if isinstance(result.get("articles"), list) else []

    inserted_jobs = supabase_rest_request(
        "POST",
        "analysis_jobs",
        json_body=[
            {
                "user_id": user_id,
                "product_id": product_id,
                "source_type": source_type,
                "source_url": source_url,
                "source_file_name": source_file_name,
                "source_mime_type": source_mime_type,
                "query": query,
                "status": "completed",
                "headline": headline,
                "summary": summary,
                "raw_result": result,
                "completed_at": datetime.now(timezone.utc).isoformat(),
            }
        ],
    )
    if not isinstance(inserted_jobs, list) or not inserted_jobs:
        raise RuntimeError("Supabase did not return the created analysis job.")

    analysis_job_id = inserted_jobs[0].get("id")
    if not analysis_job_id:
        raise RuntimeError("Supabase did not return an analysis job id.")

    topic_rows = []
    selected_lookup = {topic.lower() for topic in selected_topics}
    for index, topic in enumerate(topics):
        topic_text = str(topic).strip()
        if not topic_text:
            continue
        topic_rows.append(
            {
                "analysis_job_id": analysis_job_id,
                "topic": topic_text,
                "sort_order": index,
                "selected": topic_text.lower() in selected_lookup,
            }
        )

    inserted_topics: list[dict[str, object]] = []
    if topic_rows:
        rows = supabase_rest_request("POST", "analysis_topics", json_body=topic_rows)
        if isinstance(rows, list):
            inserted_topics = rows

    topic_id_by_name = {
        str(item.get("topic")).lower(): item.get("id")
        for item in inserted_topics
        if item.get("topic") and item.get("id")
    }

    article_rows = []
    for index, article in enumerate(articles):
        if not isinstance(article, dict):
            continue
        topic = str(article.get("topic") or "").strip()
        content = str(article.get("content") or "").strip()
        if not content:
            continue
        article_rows.append(
            {
                "analysis_job_id": analysis_job_id,
                "topic_id": topic_id_by_name.get(topic.lower()) if topic else None,
                "title": headline,
                "topic": topic or "Main topic",
                "content": content,
                "image_url": article.get("image_url"),
                "sort_order": index,
            }
        )

    if article_rows:
        supabase_rest_request("POST", "generated_articles", json_body=article_rows)


def build_missing_dependency_message(package_name: str, extra_hint: Optional[str] = None) -> str:
    message = (
        f"{package_name} is not available in the Python interpreter running this server: {sys.executable}. "
        "Start the app with .\\.venv\\Scripts\\python.exe or activate the virtual environment first."
    )
    if extra_hint:
        message = f"{message} {extra_hint}"
    return message


def get_proxy_url(scheme: str = "https") -> Optional[str]:
    if scheme == "https":
        return YOUTUBE_PROXY_HTTPS or YOUTUBE_PROXY_URL or YOUTUBE_PROXY_HTTP or None
    return YOUTUBE_PROXY_HTTP or YOUTUBE_PROXY_URL or YOUTUBE_PROXY_HTTPS or None


def remote_media_fallback_available() -> bool:
    return ENABLE_REMOTE_MEDIA_DOWNLOAD_FALLBACK or bool(APIFY_TOKEN or get_proxy_url("http") or get_proxy_url("https"))


def ensure_ffmpeg() -> None:
    if shutil.which("ffmpeg") is None:
        raise RuntimeError(
            "ffmpeg is not available on PATH. Install ffmpeg and restart the server."
        )


def get_genai_client():
    global GENAI_CLIENT
    if not GEMINI_API_KEY:
        raise RuntimeError("Gemini is not configured")
    if GENAI_CLIENT is None:
        genai = get_genai_module()
        GENAI_CLIENT = genai.Client(api_key=GEMINI_API_KEY)
    return GENAI_CLIENT


def parse_retry_delay_seconds(message: str) -> int:
    match = re.search(r"retry in\s+(\d+)", message, re.IGNORECASE)
    if match:
        return max(int(match.group(1)), 30)
    return 60


def should_skip_gemini() -> bool:
    return GEMINI_BACKOFF_UNTIL > time.time()


def gemini_generate_text(prompt: str) -> str:
    global GEMINI_BACKOFF_UNTIL
    if should_skip_gemini():
        raise RuntimeError("Gemini is temporarily cooling down after a quota/rate-limit response.")

    client = get_genai_client()
    try:
        response = client.models.generate_content(
            model=os.getenv("GEMINI_MODEL", "gemini-2.5-flash"),
            contents=prompt,
        )
        return getattr(response, "text", "").strip()
    except Exception as exc:
        error_text = str(exc)
        if "429" in error_text or "RESOURCE_EXHAUSTED" in error_text or "rate limit" in error_text.lower():
            GEMINI_BACKOFF_UNTIL = time.time() + parse_retry_delay_seconds(error_text)
        raise


def get_gemini_model_name() -> str:
    return os.getenv("GEMINI_MODEL", "gemini-2.5-flash")


def load_history() -> list[dict]:
    if not HISTORY_FILE.exists():
        return []
    try:
        with HISTORY_FILE.open("r", encoding="utf-8") as handle:
            data = json.load(handle)
            return data if isinstance(data, list) else []
    except Exception:
        return []


def save_history(items: list[dict]) -> None:
    try:
        HISTORY_FILE.parent.mkdir(parents=True, exist_ok=True)
        with HISTORY_FILE.open("w", encoding="utf-8") as handle:
            json.dump(items, handle, indent=2)
    except OSError:
        # Read-only filesystems such as serverless runtimes should not crash the app.
        return None


def add_history_entry(result: dict[str, str], source: str) -> None:
    history = load_history()
    entry = {
        "id": int(time.time() * 1000),
        "source": source,
        "headline": result.get("headline", "Media summary"),
        "summary": result.get("summary", ""),
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    history.insert(0, entry)
    save_history(history[:10])


def get_whisper_model():
    global WHISPER_MODEL
    if WHISPER_MODEL is None:
        whisper = get_whisper_module()
        WHISPER_CACHE_DIR.mkdir(parents=True, exist_ok=True)
        WHISPER_MODEL = whisper.load_model("tiny", download_root=str(WHISPER_CACHE_DIR))
    return WHISPER_MODEL


def get_embedder():
    global EMBEDDER
    if EMBEDDER is None:
        sentence_transformer = get_sentence_transformer_class()
        EMBEDDER = sentence_transformer("all-MiniLM-L6-v2")
    return EMBEDDER


def maybe_write_youtube_cookies_file(output_dir: Path) -> Optional[Path]:
    if not YOUTUBE_COOKIES_B64:
        return None
    try:
        decoded = base64.b64decode(YOUTUBE_COOKIES_B64).decode("utf-8")
    except Exception as exc:
        raise RuntimeError("YOUTUBE_COOKIES_B64 is not valid base64-encoded Netscape cookie data.") from exc

    cookies_path = output_dir / "youtube-cookies.txt"
    cookies_path.write_text(decoded, encoding="utf-8")
    return cookies_path


def cache_is_fresh(path: Path, ttl_seconds: int = SOURCE_CACHE_TTL_SECONDS) -> bool:
    try:
        return path.exists() and (time.time() - path.stat().st_mtime) <= ttl_seconds
    except OSError:
        return False


def write_json_cache(path: Path, payload: object) -> None:
    try:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
    except OSError:
        return None


def read_json_cache(path: Path) -> Optional[object]:
    if not cache_is_fresh(path):
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None


def build_content_cache_key(prefix: str, value: str) -> str:
    digest = hashlib.sha256((value or "").encode("utf-8")).hexdigest()
    return f"{prefix}-{digest}"


def source_cache_path(cache_key: str) -> Path:
    return SOURCE_CACHE_DIR / f"{cache_key}.json"


def load_cached_source_content(cache_key: str) -> Optional[dict[str, object]]:
    cached = read_json_cache(source_cache_path(cache_key))
    return cached if isinstance(cached, dict) else None


def save_cached_source_content(cache_key: str, payload: dict[str, object]) -> None:
    write_json_cache(source_cache_path(cache_key), payload)


def build_summary_cache_key(text: str, query: str) -> str:
    return build_content_cache_key("summary", f"{query}\n{text}")


def summary_cache_path(cache_key: str) -> Path:
    return SUMMARY_CACHE_DIR / f"{cache_key}.txt"


def load_cached_summary(cache_key: str) -> Optional[str]:
    path = summary_cache_path(cache_key)
    if not cache_is_fresh(path):
        return None
    try:
        return path.read_text(encoding="utf-8").strip() or None
    except OSError:
        return None


def save_cached_summary(cache_key: str, summary_text: str) -> None:
    try:
        path = summary_cache_path(cache_key)
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(summary_text, encoding="utf-8")
    except OSError:
        return None


def analysis_cache_path(cache_key: str) -> Path:
    return ANALYSIS_CACHE_DIR / f"{cache_key}.json"


def load_cached_analysis_result(cache_key: str) -> Optional[dict[str, object]]:
    cached = read_json_cache(analysis_cache_path(cache_key))
    return cached if isinstance(cached, dict) else None


def save_cached_analysis_result(cache_key: str, payload: dict[str, object]) -> None:
    write_json_cache(analysis_cache_path(cache_key), payload)


def article_cache_path(cache_key: str) -> Path:
    return ARTICLE_CACHE_DIR / f"{cache_key}.json"


def load_cached_article_result(cache_key: str) -> Optional[dict[str, object]]:
    cached = read_json_cache(article_cache_path(cache_key))
    return cached if isinstance(cached, dict) else None


def save_cached_article_result(cache_key: str, payload: dict[str, object]) -> None:
    write_json_cache(article_cache_path(cache_key), payload)


def run_with_timeout(label: str, timeout_seconds: int, func, *args, **kwargs):
    executor = ThreadPoolExecutor(max_workers=1)
    future = executor.submit(func, *args, **kwargs)
    try:
        return future.result(timeout=timeout_seconds)
    except Exception as exc:
        from concurrent.futures import TimeoutError as FutureTimeoutError
        if isinstance(exc, FutureTimeoutError):
            logger.warning("%s | timeout after %ss", label, timeout_seconds)
            future.cancel()
            executor.shutdown(wait=False, cancel_futures=True)
            raise RuntimeError(f"{label} timed out.") from exc
        raise
    finally:
        executor.shutdown(wait=False, cancel_futures=False)


def build_transcript_api_instance():
    transcript_api_class = get_youtube_transcript_api_class()
    http_proxy = get_proxy_url("http")
    https_proxy = get_proxy_url("https")
    if http_proxy or https_proxy:
        proxies_module = importlib.import_module("youtube_transcript_api.proxies")
        generic_proxy_config = getattr(proxies_module, "GenericProxyConfig")
        try:
            return transcript_api_class(
                proxy_config=generic_proxy_config(
                    http_url=http_proxy,
                    https_url=https_proxy,
                )
            )
        except TypeError:
            # Supports lightweight test doubles that do not accept proxy args.
            return transcript_api_class()
    return transcript_api_class()


def is_youtube_url(url: str) -> bool:
    host = (urlparse(url).hostname or "").lower()
    return "youtube.com" in host or "youtu.be" in host


def extract_youtube_video_id(url: str) -> Optional[str]:
    parsed = urlparse(url)
    host = (parsed.hostname or "").lower()
    if "youtu.be" in host:
        candidate = parsed.path.strip("/").split("/")[0]
        return candidate or None
    if "youtube.com" in host:
        if parsed.path == "/watch":
            for pair in parsed.query.split("&"):
                if pair.startswith("v="):
                    return pair.split("=", 1)[1] or None
        path_parts = [part for part in parsed.path.split("/") if part]
        if len(path_parts) >= 2 and path_parts[0] in {"shorts", "embed", "live"}:
            return path_parts[1]
    return None


def fetch_youtube_transcript_text(url: str) -> str:
    video_id = extract_youtube_video_id(url)
    if not video_id:
        raise RuntimeError("Could not extract a valid YouTube video id from the URL.")

    transcript_api = build_transcript_api_instance()
    try:
        transcript_entries = transcript_api.fetch(video_id, languages=["en", "en-US", "en-GB"])
    except Exception:
        transcript_entries = transcript_api.fetch(video_id)

    transcript_text = " ".join(
        str(getattr(item, "text", "")).replace("\n", " ").strip()
        for item in transcript_entries
        if getattr(item, "text", "")
    ).strip()
    if not transcript_text:
        raise RuntimeError("No transcript text was available for this YouTube video.")
    return transcript_text


def fetch_youtube_metadata(url: str) -> dict[str, object]:
    yt_dlp = get_yt_dlp_module()
    output_dir = Path(tempfile.mkdtemp(prefix="yt-meta-"))
    cookies_file = maybe_write_youtube_cookies_file(output_dir)
    opts: dict[str, object] = {
        "skip_download": True,
        "quiet": True,
        "no_warnings": True,
        "noprogress": True,
        "socket_timeout": 25,
        "retries": 2,
        "logger": QuietYtdlpLogger(),
        "http_headers": {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/127.0.0.0 Safari/537.36"
            )
        },
    }
    proxy_url = get_proxy_url("https")
    if proxy_url:
        opts["proxy"] = proxy_url
    if cookies_file:
        opts["cookiefile"] = str(cookies_file)
    try:
        with yt_dlp.YoutubeDL(opts) as ydl:
            info = normalize_yt_info_entry(ydl.extract_info(url, download=False))
    finally:
        shutil.rmtree(output_dir, ignore_errors=True)
    return {
        "title": str(info.get("title") or "").strip(),
        "description": str(info.get("description") or "").strip(),
        "channel": str(info.get("channel") or info.get("uploader") or "").strip(),
        "duration": int(info.get("duration") or 0),
        "categories": [str(item).strip() for item in (info.get("categories") or []) if str(item).strip()],
        "tags": [str(item).strip() for item in (info.get("tags") or [])[:12] if str(item).strip()],
        "webpage_url": str(info.get("webpage_url") or url).strip(),
    }


def sanitize_youtube_metadata_text(text: str) -> str:
    cleaned = unescape(text or "")
    cleaned = re.sub(r"https?://\S+", " ", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"www\.\S+", " ", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"@\w+", " ", cleaned)
    cleaned = re.sub(r"#\w+", " ", cleaned)
    cleaned = re.sub(r"\b(?:instagram|youtube|facebook|twitter|tiktok)\b", " ", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\b(?:follow|subscribe|credits?|link in bio)\b", " ", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip()


def build_youtube_source_text(metadata: dict[str, object], transcript_text: str = "") -> str:
    title = sanitize_youtube_metadata_text(str(metadata.get("title") or ""))
    transcript_text = clean_source_text(transcript_text or "")
    if transcript_text:
        pieces = [
            f"Video title: {title}" if title else "",
            transcript_text,
        ]
        return clean_source_text("\n".join(piece for piece in pieces if piece).strip())

    description = sanitize_youtube_metadata_text(str(metadata.get("description") or ""))
    pieces = [
        f"Video title: {title}" if title else "",
        f"Short description: {description[:280]}" if description else "",
    ]
    return clean_source_text("\n".join(piece for piece in pieces if piece).strip())


def clean_vtt_caption_text(raw_text: str) -> str:
    cleaned_lines: list[str] = []
    seen_recent: str = ""
    for raw_line in raw_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line == "WEBVTT" or line.startswith(("Kind:", "Language:")):
            continue
        if "-->" in line:
            continue
        line = re.sub(r"<\d{2}:\d{2}:\d{2}\.\d{3}>", " ", line)
        line = re.sub(r"</?c[^>]*>", "", line)
        line = re.sub(r"</?[^>]+>", "", line)
        line = " ".join(line.split()).strip()
        if not line:
            continue
        if line == seen_recent:
            continue
        cleaned_lines.append(line)
        seen_recent = line
    return " ".join(cleaned_lines).strip()


def convert_asterisk_bold_to_html(text: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text or "")


def strip_html_tags(text: str) -> str:
    no_tags = re.sub(r"<[^>]+>", " ", text or "")
    return " ".join(no_tags.split()).strip()


def sanitize_article_html(raw_html: str) -> str:
    if not raw_html.strip():
        return ""

    html = convert_asterisk_bold_to_html(raw_html.strip())
    html = re.sub(r"(?m)^\s*#{3}\s*(.+)$", r"<h3>\1</h3>", html)
    html = re.sub(r"(?m)^\s*#{2}\s*(.+)$", r"<h2>\1</h2>", html)

    if not re.search(r"<(?:h2|h3|p|ul|li)\b", html):
        blocks: list[str] = []
        paragraphs = [block.strip() for block in re.split(r"\n\s*\n", html) if block.strip()]
        for index, paragraph in enumerate(paragraphs):
            cleaned = paragraph.strip()
            if cleaned.startswith("- "):
                items = [item.strip()[2:].strip() for item in cleaned.splitlines() if item.strip().startswith("- ")]
                blocks.append("<ul>" + "".join(f"<li>{item}</li>" for item in items if item) + "</ul>")
            elif index == 0:
                blocks.append(f"<h2>{cleaned}</h2>")
            else:
                blocks.append(f"<p>{cleaned}</p>")
        html = "\n".join(blocks)

    html = re.sub(r"\n{2,}", "\n", html)
    html = re.sub(r"</?(script|style)[^>]*>", "", html, flags=re.IGNORECASE)
    html = re.sub(
        r"</?([a-zA-Z0-9]+)(?:\s+[^>]*)?>",
        lambda match: match.group(0) if match.group(1).lower() in ALLOWED_ARTICLE_TAGS else "",
        html,
    )
    normalized_lines: list[str] = []
    for raw_line in re.split(r"\n+", html):
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("<"):
            normalized_lines.append(line)
        else:
            normalized_lines.append(f"<p>{line}</p>")
    return "\n".join(normalized_lines).strip()


def sanitize_inline_html(text: str) -> str:
    html = convert_asterisk_bold_to_html(text or "")
    html = re.sub(r"</?(script|style)[^>]*>", "", html, flags=re.IGNORECASE)
    html = re.sub(
        r"</?([a-zA-Z0-9]+)(?:\s+[^>]*)?>",
        lambda match: match.group(0) if match.group(1).lower() == "strong" else "",
        html,
    )
    return html.strip()


def remove_timestamps(text: str) -> str:
    cleaned = re.sub(r"\[\s*\d{1,2}:\d{2}(?::\d{2})?\s*\]", " ", text or "")
    cleaned = re.sub(r"\(\s*\d{1,2}:\d{2}(?::\d{2})?\s*\)", " ", cleaned)
    cleaned = re.sub(r"\b\d{1,2}:\d{2}(?::\d{2})?\b", " ", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip()


def contains_metadata_pollution(text: str) -> bool:
    candidate = (text or "").lower()
    if not candidate:
        return False
    patterns = [
        r"https?://",
        r"www\.",
        r"instagram\.com",
        r"youtube\.com",
        r"\binstagram\b",
        r"\bcredits?\b",
        r"\bfollow\b",
        r"\bsubscribe\b",
        r"\blink in bio\b",
    ]
    return any(re.search(pattern, candidate) for pattern in patterns)


def analysis_output_is_valid(result: dict[str, object]) -> bool:
    heading = str(result.get("heading") or result.get("headline") or "")
    summary = str(result.get("summary") or "")
    key_points = result.get("key_points", [])
    topics = result.get("topics", [])

    text_fragments = [heading, summary]
    if isinstance(key_points, list):
        text_fragments.extend(str(item) for item in key_points)
    if isinstance(topics, list):
        for topic in topics:
            if not isinstance(topic, dict):
                return False
            title = clean_topic_title(str(topic.get("title", "")))
            if not title or title.lower() in SUSPICIOUS_ANALYSIS_TITLES:
                return False
            text_fragments.append(title)
            text_fragments.append(str(topic.get("explanation", "")))
            text_fragments.append(str(topic.get("importance", "")))
            points = topic.get("points", [])
            if not isinstance(points, list) or not points:
                return False
            for point in points:
                if not isinstance(point, dict):
                    return False
                text_fragments.append(str(point.get("label", "")))
                text_fragments.append(str(point.get("description", "")))

    normalized_fragments = [remove_timestamps(strip_html_tags(fragment)).strip() for fragment in text_fragments if str(fragment).strip()]
    if any(contains_metadata_pollution(fragment) for fragment in normalized_fragments):
        return False
    unique_fragments = {re.sub(r"\W+", " ", fragment.lower()).strip() for fragment in normalized_fragments if fragment}
    if len(unique_fragments) < max(4, len(normalized_fragments) // 3):
        return False
    return bool(heading.strip() and summary.strip() and isinstance(topics, list) and topics)


def direct_gemini_failure_reason(exc: Exception) -> str:
    message = str(exc or "").strip() or exc.__class__.__name__
    lower = message.lower()
    if "unsupported" in lower or "not supported" in lower:
        return "The current Gemini request format does not support direct video URL analysis."
    if "timeout" in lower:
        return "The direct Gemini analysis request timed out."
    if "could not analyze" in lower or "access" in lower:
        return "Gemini could not access enough video content from the provided URL."
    return message


def clean_source_text(text: str) -> str:
    normalized = unescape((text or "").replace("\r", "\n"))
    normalized = re.sub(r"<[^>]+>", " ", normalized)
    normalized = re.sub(r"\b\d{1,2}:\d{2}(?::\d{2})?\b", " ", normalized)
    normalized = normalized.replace("♪", " ")
    cleaned_lines: list[str] = []
    seen = set()
    for raw_line in normalized.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        line = re.sub(r"^[A-Z][A-Za-z0-9 _-]{0,24}:\s*", "", line)
        for pattern in FILLER_PATTERNS:
            line = re.sub(pattern, " ", line, flags=re.IGNORECASE)
        line = re.sub(r"\s{2,}", " ", line).strip(" -")
        if len(line) < 20:
            continue
        lowered = line.lower()
        if lowered in seen:
            continue
        seen.add(lowered)
        cleaned_lines.append(line)
    joined = " ".join(cleaned_lines)
    sentences = split_sentences(joined)
    deduped_sentences: list[str] = []
    seen_sentences = set()
    for sentence in sentences:
        normalized_sentence = re.sub(r"\W+", " ", sentence.lower()).strip()
        if not normalized_sentence or normalized_sentence in seen_sentences:
            continue
        seen_sentences.add(normalized_sentence)
        deduped_sentences.append(sentence)
    return " ".join(deduped_sentences).strip()


def extract_webpage_content(url: str) -> dict[str, object]:
    cache_key = build_content_cache_key("webpage", url)
    cached = load_cached_source_content(cache_key)
    if cached:
        logger.info("webpage-cache-hit | %s", url)
        return cached

    BeautifulSoup = get_beautiful_soup_class()
    try:
        response = httpx.get(
            url,
            timeout=20.0,
            follow_redirects=True,
            headers={
                "User-Agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/127.0.0.0 Safari/537.36"
                )
            },
        )
        response.raise_for_status()
    except httpx.HTTPStatusError as exc:
        status_code = exc.response.status_code if exc.response is not None else None
        if status_code in {401, 403}:
            raise RuntimeError("That website blocked direct content extraction. Please try another public URL or upload the media file.") from exc
        if status_code == 404:
            raise RuntimeError("That URL could not be found. Please check the link and try again.") from exc
        raise RuntimeError("We could not fetch readable content from that URL right now. Please try another source.") from exc
    except httpx.HTTPError as exc:
        raise RuntimeError("We could not fetch readable content from that URL right now. Please try another source.") from exc
    content_type = response.headers.get("content-type", "").lower()
    if "text/html" not in content_type and "application/xhtml+xml" not in content_type:
        raise RuntimeError("Only webpage URLs are supported for direct link analysis right now.")

    soup = BeautifulSoup(response.text, "html.parser")
    for tag_name in ("script", "style", "noscript", "svg", "footer", "nav", "aside"):
        for tag in soup.find_all(tag_name):
            tag.decompose()

    title = ""
    if soup.title and soup.title.string:
        title = soup.title.string.strip()
    meta_description = ""
    meta_tag = soup.find("meta", attrs={"name": "description"}) or soup.find("meta", attrs={"property": "og:description"})
    if meta_tag and meta_tag.get("content"):
        meta_description = str(meta_tag.get("content")).strip()

    headings = [
        " ".join(tag.get_text(" ", strip=True).split())
        for tag in soup.find_all(["h1", "h2", "h3"], limit=12)
        if tag.get_text(" ", strip=True)
    ]
    paragraphs = [
        " ".join(tag.get_text(" ", strip=True).split())
        for tag in soup.find_all(["p", "li"], limit=80)
        if tag.get_text(" ", strip=True)
    ]
    raw_content = "\n".join(filter(None, [title, meta_description, *headings, *paragraphs]))
    cleaned = clean_source_text(raw_content)
    if len(cleaned) < 120:
        raise RuntimeError("We could not extract enough readable text from that URL.")

    payload = {
        "title": title or urlparse(url).netloc,
        "meta_description": meta_description,
        "headings": headings,
        "content": cleaned,
    }
    save_cached_source_content(cache_key, payload)
    return payload


def sanitize_topic_point(point: dict[str, object]) -> Optional[dict[str, str]]:
    if not isinstance(point, dict):
        return None
    label = sanitize_inline_html(str(point.get("label", "")).strip())
    description = sanitize_inline_html(str(point.get("description", "")).strip())
    label = remove_timestamps(label)
    description = remove_timestamps(description)
    if not label and not description:
        return None
    return {
        "label": label or "Key idea",
        "description": description,
    }


def build_topic_details_bundle(summary_text: str, structured_topics: Optional[list[dict[str, object]]] = None) -> tuple[list[dict[str, object]], bool]:
    raw_details = structured_topics or get_topic_details_from_summary(summary_text)
    details: list[dict[str, str]] = []
    seen_titles: set[str] = set()
    for item in raw_details:
        title = clean_topic_title(str(item.get("title", "")))
        key = title.lower()
        if not title or key in seen_titles:
            continue
        seen_titles.add(key)
        points = []
        for raw_point in item.get("points", []) if isinstance(item, dict) else []:
            cleaned_point = sanitize_topic_point(raw_point)
            if cleaned_point:
                points.append(cleaned_point)
        explanation = remove_timestamps(sanitize_inline_html(str(item.get("explanation", ""))))
        importance = remove_timestamps(sanitize_inline_html(str(item.get("importance", ""))))
        details.append(
            {
                "id": f"topic_{len(details) + 1}",
                "title": title,
                "summary": explanation,
                "explanation": explanation,
                "importance": importance,
                "points": points,
            }
        )
    used_fallback = all(title_needs_editorial_rewrite(item.get("title", "")) for item in details[: min(3, len(details))]) if details else True
    return details, used_fallback


def build_article_slug(title: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", strip_html_tags(title).lower()).strip("-")
    return slug[:80] or "generated-article"


def build_meta_title(headline: str, topic: str, article_type: str) -> str:
    base = strip_html_tags(topic or headline or article_type)
    fallback = strip_html_tags(headline or article_type)
    candidate = base if len(base) >= 18 else f"{base} | {fallback}".strip(" |")
    return candidate[:60].strip() or fallback[:60].strip() or "Generated Article"


def build_meta_description(summary_text: str) -> str:
    summary_points = [line[2:].strip() for line in summary_text.splitlines() if line.strip().startswith("-")]
    source = " ".join(summary_points) if summary_points else summary_text
    source = re.sub(r"\s+", " ", strip_html_tags(source)).strip()
    return source[:157].rstrip(" ,.;:-") + ("..." if len(source) > 157 else "")


def build_geo_keywords(topic: str, secondary_keywords: list[str]) -> list[str]:
    seed = [topic] + secondary_keywords[:3]
    phrases: list[str] = []
    for keyword in seed:
        clean = strip_html_tags(keyword).strip()
        if not clean:
            continue
        phrases.extend([
            f"What is {clean}?",
            f"Why does {clean} matter?",
            f"{clean} explained",
        ])
    return list(dict.fromkeys(phrases))[:6]


def build_dynamic_subheadings(topic: str, summary_points: list[str], article_type: str) -> list[str]:
    clean_topic = strip_html_tags(topic or "Main Topic").strip() or "Main Topic"
    keywords = extract_top_keywords(" ".join(summary_points) or clean_topic, limit=4)
    lead_keyword = keywords[0] if keywords else clean_topic
    support_keyword = keywords[1] if len(keywords) > 1 else clean_topic
    context_keyword = keywords[2] if len(keywords) > 2 else support_keyword

    if article_type == "News Article":
        return [
            f"Key Developments Around {clean_topic}",
            f"The Context Behind {lead_keyword}",
            f"Why {support_keyword} Matters Now",
            f"What to Watch Next in {context_keyword}",
        ]
    if article_type == "How-to Guide":
        return [
            f"Understanding the Core Challenge in {clean_topic}",
            f"Steps That Move {lead_keyword} Forward",
            f"Common Mistakes Around {support_keyword}",
            f"How to Apply These Lessons in Practice",
        ]
    if article_type == "Listicle":
        return [
            f"The First Shift Shaping {clean_topic}",
            f"Why {lead_keyword} Changes the Conversation",
            f"Lessons Hidden Inside {support_keyword}",
            f"What Makes {context_keyword} Worth Watching",
        ]
    if article_type == "Review":
        return [
            f"What Holds {clean_topic} Together",
            f"Where {lead_keyword} Delivers",
            f"Where {support_keyword} Falls Short",
            f"The Larger Takeaway for Readers",
        ]
    if article_type == "SEO Article":
        return [
            f"What {clean_topic} Means for Search Intent",
            f"The Core Signals Behind {lead_keyword}",
            f"How {support_keyword} Shapes the Topic",
            f"Why {context_keyword} Deserves Closer Attention",
        ]
    return [
        f"How {clean_topic} Takes Shape",
        f"What {lead_keyword} Reveals",
        f"Why {support_keyword} Carries Weight",
        f"Where the Bigger Argument Lands",
    ]


def deduplicate_article_html(article_html: str) -> str:
    blocks = re.findall(r"<(h2|h3|p|li)>(.*?)</\1>", article_html, flags=re.IGNORECASE | re.DOTALL)
    seen_blocks = set()
    cleaned_chunks: list[str] = []
    for tag, inner in blocks:
        text = re.sub(r"\s+", " ", strip_html_tags(inner)).strip()
        normalized = re.sub(r"\W+", " ", text.lower()).strip()
        if not normalized or normalized in seen_blocks:
            continue
        seen_blocks.add(normalized)
        if tag.lower() == "p":
            sentences = split_sentences(text)
            unique_sentences: list[str] = []
            seen_sentences = set()
            for sentence in sentences:
                ns = re.sub(r"\W+", " ", sentence.lower()).strip()
                if not ns or ns in seen_sentences:
                    continue
                seen_sentences.add(ns)
                unique_sentences.append(sentence)
            text = " ".join(unique_sentences) or text
            cleaned_chunks.append(f"<p>{text}</p>")
        else:
            cleaned_chunks.append(f"<{tag.lower()}>{inner}</{tag.lower()}>")
    return postprocess_article_html("\n".join(cleaned_chunks))


def score_article_output(
    *,
    title: str,
    meta_title: str,
    meta_description: str,
    focus_keyword: str,
    secondary_keywords: list[str],
    geo_keywords: list[str],
    content_html: str,
    article_type: str,
) -> dict[str, object]:
    text = strip_html_tags(content_html)
    h2_count = len(re.findall(r"<h2>", content_html, flags=re.IGNORECASE))
    h3_count = len(re.findall(r"<h3>", content_html, flags=re.IGNORECASE))
    p_count = len(re.findall(r"<p>", content_html, flags=re.IGNORECASE))
    seo_score = 0
    geo_score = 0
    seo_checks = []
    geo_checks = []

    title_ok = 20 <= len(title) <= 70
    seo_score += 2 if title_ok else 1
    seo_checks.append({"label": "Title quality", "score": 2 if title_ok else 1, "note": "Title length and clarity."})

    meta_ok = 35 <= len(meta_title) <= 60 and 80 <= len(meta_description) <= 160
    seo_score += 2 if meta_ok else 1
    seo_checks.append({"label": "Meta description", "score": 2 if meta_ok else 1, "note": "Meta title and description are within useful SEO ranges."})

    heading_ok = h2_count >= 1 and h3_count >= 2
    seo_score += 2 if heading_ok else 1
    seo_checks.append({"label": "Heading structure", "score": 2 if heading_ok else 1, "note": "The article uses a readable heading hierarchy."})

    keyword_mentions = sum(1 for kw in [focus_keyword, *secondary_keywords[:3]] if kw and kw.lower() in text.lower())
    seo_score += 2 if keyword_mentions >= 2 else 1
    seo_checks.append({"label": "Keyword usage", "score": 2 if keyword_mentions >= 2 else 1, "note": "Primary and secondary keywords are reflected naturally."})

    readability_ok = p_count >= 5 and len(text.split()) >= 450
    seo_score += 2 if readability_ok else 1
    seo_checks.append({"label": "Readability", "score": 2 if readability_ok else 1, "note": "Paragraph depth and article length support readability."})

    search_intent_ok = article_type in ARTICLE_TYPE_INSTRUCTIONS and focus_keyword.lower() in text.lower()
    geo_score += 3 if search_intent_ok else 1
    geo_checks.append({"label": "Search intent match", "score": 3 if search_intent_ok else 1, "note": "The article aligns with the selected format and core topic."})

    geo_density = sum(1 for phrase in geo_keywords[:4] if phrase.split("?")[0].lower() in text.lower())
    geo_score += 3 if geo_density >= 1 else 1
    geo_checks.append({"label": "AI search phrasing", "score": 3 if geo_density >= 1 else 1, "note": "The article includes answer-style phrasing for AI search engines."})

    uniqueness_ok = len(set(split_sentences(text))) >= max(8, len(split_sentences(text)) // 2)
    geo_score += 2 if uniqueness_ok else 1
    geo_checks.append({"label": "Content uniqueness", "score": 2 if uniqueness_ok else 1, "note": "Repeated sections are minimized."})

    local_relevance_ok = any(term.lower() in text.lower() for term in ["market", "audience", "readers", "region", "industry"])
    geo_score += 2 if local_relevance_ok else 1
    geo_checks.append({"label": "Context relevance", "score": 2 if local_relevance_ok else 1, "note": "The article reflects who the content is for and why it matters."})

    suggestions: list[str] = []
    if not meta_ok:
        suggestions.append("Tighten the meta title and meta description so they fit standard SEO lengths.")
    if not heading_ok:
        suggestions.append("Add more meaningful section headings to improve scanability.")
    if keyword_mentions < 2:
        suggestions.append("Weave the focus keyword and supporting terms into the article more naturally.")
    if geo_density < 1:
        suggestions.append("Add answer-style phrases that align with AI search and summary engines.")

    return {
        "seoScore": min(seo_score, 10),
        "geoScore": min(geo_score, 10),
        "seoChecks": seo_checks,
        "geoChecks": geo_checks,
        "improvementSuggestions": suggestions,
    }


def build_article_package(
    *,
    headline_text: str,
    summary_text: str,
    topic: str,
    topic_summary: str = "",
    topic_points: list[dict[str, str]] | None = None,
    article_type: str,
    source_context: str,
    target_audience: str,
    variant_index: int = 0,
) -> dict[str, object]:
    profiler = StageProfiler(f"article-package:{clean_topic_title(topic or headline_text)}#{variant_index + 1}")
    with profiler.stage("article_pipeline"):
        content = generate_article_html(
            headline_text=headline_text,
            summary_text=summary_text,
            topic=topic,
            topic_summary=topic_summary,
            topic_points=topic_points or [],
            article_type=article_type,
            source_context=source_context,
            target_audience=target_audience,
            variant_index=variant_index,
            profiler=profiler,
        )
    focus_keyword = strip_html_tags(topic).strip() or strip_html_tags(headline_text).strip()
    secondary_keywords = [
        keyword for keyword in extract_top_keywords(f"{summary_text} {topic}", limit=6)
        if keyword.lower() != focus_keyword.lower()
    ][:5]
    geo_keywords = build_geo_keywords(focus_keyword, secondary_keywords)
    meta_title = build_meta_title(headline_text, topic, article_type)
    meta_description = build_meta_description(summary_text)
    with profiler.stage("seo_geo_scoring"):
        scoring = score_article_output(
            title=headline_text,
            meta_title=meta_title,
            meta_description=meta_description,
            focus_keyword=focus_keyword,
            secondary_keywords=secondary_keywords,
            geo_keywords=geo_keywords,
            content_html=content,
            article_type=article_type,
        )
    payload = {
        "topic": topic,
        "article_type": article_type,
        "content": content,
        "image_url": build_article_image_url(topic),
        "meta_title": meta_title,
        "meta_description": meta_description,
        "slug": build_article_slug(meta_title),
        "focus_keyword": focus_keyword,
        "secondary_keywords": secondary_keywords,
        "geo_keywords": geo_keywords,
        "seo_report": scoring,
    }
    profiler.log_total()
    return payload


def clean_topic_title(title: str) -> str:
    cleaned = strip_html_tags(convert_asterisk_bold_to_html(title or ""))
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" -,:;")
    return cleaned[:90].strip() or "Main Topic"


def title_needs_editorial_rewrite(title: str) -> bool:
    cleaned = clean_topic_title(title)
    lower = cleaned.lower()
    if lower in GENERIC_TOPIC_TITLES:
        return True
    if len(cleaned.split()) < 2:
        return True
    if re.fullmatch(r"[A-Za-z]+", cleaned) and len(cleaned) < 12:
        return True
    return False


def build_editorial_topic_title(keyword: str, support_text: str, index: int) -> str:
    support = " ".join((support_text or "").split())
    keyword_title = keyword.strip().title() or "Main Topic"
    if not support:
        return keyword_title

    support = re.sub(r"^[\"'(\[]+|[\"')\].,;:!?]+$", "", support)
    if ":" in support:
        prefix, suffix = support.split(":", 1)
        candidate = suffix.strip() if len(suffix.split()) >= 3 else prefix.strip()
        if candidate:
            support = candidate

    match = re.match(
        r"^(?P<subject>[A-Z][A-Za-z0-9&' -]{2,40})\s+(?P<verb>is|are|was|were|has|have|can|will)\s+(?P<rest>.+)$",
        support,
    )
    if match:
        subject = match.group("subject").strip(" ,.;:-")
        verb = match.group("verb").strip()
        rest = match.group("rest").strip(" ,.;:-")
        if subject.lower() not in {"the source", "the video", "the discussion", "the speaker", "the summary"}:
            return clean_topic_title(f"How {subject} {verb} {rest}")

    words = support.split()
    if len(words) >= 5:
        phrase = " ".join(words[:9]).strip(" ,.;:-")
        lower_phrase = phrase.lower()
        if keyword.lower() in lower_phrase and index > 2:
            return clean_topic_title(phrase)
        if index == 0:
            return clean_topic_title(f"{keyword_title} and the Bigger Shift")
        if index == 1:
            return clean_topic_title(f"How {keyword_title} Shapes the Debate")
        if index == 2:
            return clean_topic_title(f"What {keyword_title} Reveals")
        return clean_topic_title(f"{keyword_title}: {phrase}")

    return clean_topic_title(keyword_title)


def postprocess_article_html(raw_html: str) -> str:
    html = sanitize_article_html(raw_html)
    if not html:
        return ""

    html = re.sub(
        r"<h3>\s*(Why This Topic Stands Out|Key Developments|What It Suggests|Selected Topic|Source Material)\s*</h3>",
        "",
        html,
        flags=re.IGNORECASE,
    )
    html = re.sub(
        r"<p>\s*(This article discusses|The source material says|The selected topic is|This article stays focused on)[^<]*</p>",
        "",
        html,
        flags=re.IGNORECASE,
    )
    html = re.sub(r"\n{2,}", "\n", html).strip()
    return html


def fetch_youtube_subtitles_text(url: str, output_dir: str) -> str:
    yt_dlp = get_yt_dlp_module()
    output_dir_path = Path(output_dir)
    output_dir_path.mkdir(parents=True, exist_ok=True)
    cookies_file = maybe_write_youtube_cookies_file(output_dir_path)

    def build_subtitle_opts(proxy_url: Optional[str]) -> dict[str, object]:
        opts: dict[str, object] = {
            "skip_download": True,
            "writesubtitles": True,
            "writeautomaticsub": True,
            "subtitleslangs": ["en", "en-US", "en-GB"],
            "subtitlesformat": "vtt",
            "outtmpl": str(output_dir_path / "subtitle-download.%(ext)s"),
            "quiet": True,
            "no_warnings": True,
            "noprogress": True,
            "socket_timeout": 30,
            "retries": 2,
            "logger": QuietYtdlpLogger(),
            "http_headers": {
                "User-Agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/127.0.0.0 Safari/537.36"
                )
            },
        }
        if cookies_file:
            opts["cookiefile"] = str(cookies_file)
        if proxy_url:
            opts["proxy"] = proxy_url
        return opts

    direct_error: Optional[Exception] = None
    try:
        with yt_dlp.YoutubeDL(build_subtitle_opts(None)) as ydl:
            ydl.extract_info(url, download=True)
    except Exception as exc:
        direct_error = exc
        proxy_url = get_proxy_url("https")
        if not proxy_url:
            raise
        with yt_dlp.YoutubeDL(build_subtitle_opts(proxy_url)) as ydl:
            ydl.extract_info(url, download=True)

    subtitle_files = sorted(output_dir_path.glob("*.vtt"))
    if not subtitle_files:
        if direct_error:
            raise RuntimeError(f"No YouTube subtitle file was available for this video. Direct subtitle fetch failed: {direct_error}")
        raise RuntimeError("No YouTube subtitle file was available for this video.")

    best_text = ""
    for subtitle_file in subtitle_files:
        subtitle_text = clean_vtt_caption_text(subtitle_file.read_text(encoding="utf-8", errors="ignore"))
        if len(subtitle_text) > len(best_text):
            best_text = subtitle_text
    if not best_text:
        raise RuntimeError("YouTube subtitle download succeeded, but no readable subtitle text was found.")
    return best_text


def select_apify_download_url(item: dict[str, object]) -> Optional[str]:
    candidate_keys = (
        "downloadUrl",
        "download_url",
        "apify_storage_url",
        "storageUrl",
        "fileUrl",
        "file_url",
        "url",
    )
    for key in candidate_keys:
        value = item.get(key)
        if isinstance(value, str) and value.startswith(("http://", "https://")):
            if key == "url" and is_youtube_url(value):
                continue
            return value
    return None


def find_http_urls_in_payload(payload: object) -> list[str]:
    urls: list[str] = []

    def walk(value: object) -> None:
        if isinstance(value, str):
            if value.startswith(("http://", "https://")):
                urls.append(value)
            return
        if isinstance(value, dict):
            for nested_value in value.values():
                walk(nested_value)
            return
        if isinstance(value, list):
            for nested_value in value:
                walk(nested_value)

    walk(payload)
    return urls


def select_any_non_youtube_url(payload: object) -> Optional[str]:
    for value in find_http_urls_in_payload(payload):
        if not is_youtube_url(value):
            return value
    return None


def parse_json_response_or_raise(response: httpx.Response, context: str) -> object:
    try:
        return response.json()
    except ValueError as exc:
        text = response.text.strip()
        preview = text[:300] if text else "empty response"
        raise RuntimeError(f"{context} returned a non-JSON response: {preview}") from exc


def infer_download_suffix(download_url: str, fallback: str = ".mp3") -> str:
    path = urlparse(download_url).path
    suffix = Path(path).suffix.lower()
    return suffix if suffix else fallback


def download_file_to_path(download_url: str, destination: Path) -> Path:
    with httpx.stream(
        "GET",
        download_url,
        follow_redirects=True,
        timeout=120.0,
        proxy=get_proxy_url("https"),
    ) as response:
        if response.status_code >= 400:
            text = response.text.strip()
            preview = text[:300] if text else f"HTTP {response.status_code}"
            raise RuntimeError(f"Media download failed from upstream ({response.status_code}): {preview}")
        content_type = (response.headers.get("content-type") or "").lower()
        if "application/json" in content_type or "text/plain" in content_type:
            text = response.text.strip()
            if text and len(text) < 500 and "upstream error" in text.lower():
                raise RuntimeError(f"Media download failed from upstream: {text}")
        with destination.open("wb") as file_handle:
            for chunk in response.iter_bytes():
                file_handle.write(chunk)
    return destination


def apify_kv_record_url(store_id: str, record_key: str) -> str:
    return f"https://api.apify.com/v2/key-value-stores/{store_id}/records/{record_key}?disableRedirect=true&token={APIFY_TOKEN}"


def select_apify_kv_media_url(store_id: str) -> Optional[str]:
    response = httpx.get(
        f"https://api.apify.com/v2/key-value-stores/{store_id}/keys",
        params={"token": APIFY_TOKEN},
        timeout=30.0,
        proxy=get_proxy_url("https"),
    )
    response.raise_for_status()
    payload = parse_json_response_or_raise(response, "Apify key-value store listing")
    items = payload.get("items") if isinstance(payload, dict) else None
    if not isinstance(items, list):
        return None

    def is_media_item(item: dict[str, object]) -> bool:
        key = str(item.get("key") or "").lower()
        content_type = str(item.get("contentType") or "").lower()
        return (
            content_type.startswith("audio/")
            or content_type.startswith("video/")
            or key.endswith((".mp3", ".mp4", ".m4a", ".wav", ".webm", ".aac", ".ogg", ".opus"))
        )

    for item in items:
        if isinstance(item, dict) and is_media_item(item):
            key = item.get("key")
            if isinstance(key, str) and key:
                return apify_kv_record_url(store_id, key)
    return None


def download_audio_via_apify(url: str, output_dir: str) -> str:
    if not APIFY_TOKEN:
        raise RuntimeError("Apify token is not configured")

    apify_client_class = get_apify_client_class()
    client = apify_client_class(APIFY_TOKEN)
    run_input = {
        "videos": [{"url": url}],
        "storeInKVStore": False,
        "preferredQuality": "144p",
        "preferredFormat": "mp3",
    }
    run = client.actor("streamers/youtube-video-downloader").call(run_input=run_input)
    dataset_id = getattr(run, "default_dataset_id", None) or run.get("defaultDatasetId") or run.get("default_dataset_id")
    key_value_store_id = (
        getattr(run, "default_key_value_store_id", None)
        or run.get("defaultKeyValueStoreId")
        or run.get("default_key_value_store_id")
    )
    if not dataset_id:
        raise RuntimeError("Apify actor finished without a dataset id.")

    items = list(client.dataset(dataset_id).iterate_items())
    if not items:
        raise RuntimeError("Apify actor did not return any downloadable items.")

    first_item = items[0]
    if not isinstance(first_item, dict):
        raise RuntimeError("Apify actor returned an invalid item payload.")

    status = str(first_item.get("status") or "").lower()
    if "fail" in status or "error" in status:
        raise RuntimeError(str(first_item.get("error") or first_item.get("message") or "Apify download failed."))

    download_url = select_apify_download_url(first_item)
    if not download_url:
        download_url = select_any_non_youtube_url(first_item)
    if not download_url and key_value_store_id:
        download_url = select_apify_kv_media_url(str(key_value_store_id))
    if not download_url:
        raise RuntimeError(
            "Apify actor completed, but no downloadable file URL was returned from the dataset or key-value store."
        )

    output_dir_path = Path(output_dir)
    output_dir_path.mkdir(parents=True, exist_ok=True)
    suffix = infer_download_suffix(download_url, fallback=".mp3")
    downloaded_path = output_dir_path / f"apify-download{suffix}"
    try:
        download_file_to_path(download_url, downloaded_path)
    except Exception as exc:
        raise RuntimeError(f"Apify returned a download URL, but the media fetch failed: {exc}") from exc

    if suffix in {".mp3", ".m4a", ".wav", ".aac", ".ogg", ".opus", ".webm"}:
        return str(downloaded_path)

    extracted_audio_path = output_dir_path / "apify-extracted-audio.wav"
    return extract_audio_from_video(str(downloaded_path), str(extracted_audio_path))


def normalize_yt_info_entry(info: dict[str, object]) -> dict[str, object]:
    entries = info.get("entries")
    if isinstance(entries, list):
        for entry in entries:
            if isinstance(entry, dict):
                return entry
    return info


def select_best_audio_format_id(info: dict[str, object]) -> Optional[str]:
    normalized_info = normalize_yt_info_entry(info)
    formats = normalized_info.get("formats")
    if not isinstance(formats, list):
        return None

    audio_only: list[dict[str, object]] = []
    muxed_with_audio: list[dict[str, object]] = []
    for fmt in formats:
        if not isinstance(fmt, dict):
            continue
        acodec = str(fmt.get("acodec") or "none")
        vcodec = str(fmt.get("vcodec") or "none")
        format_id = fmt.get("format_id")
        if not format_id or acodec == "none":
            continue
        if vcodec == "none":
            audio_only.append(fmt)
        else:
            muxed_with_audio.append(fmt)

    def format_rank(fmt: dict[str, object]) -> tuple[float, float, int]:
        abr = float(fmt.get("abr") or 0)
        tbr = float(fmt.get("tbr") or 0)
        ext = str(fmt.get("ext") or "")
        ext_preference = 1 if ext in {"m4a", "mp4", "webm"} else 0
        return (abr, tbr, ext_preference)

    if audio_only:
        best = max(audio_only, key=format_rank)
        return str(best.get("format_id"))
    if muxed_with_audio:
        best = max(muxed_with_audio, key=format_rank)
        return str(best.get("format_id"))
    return None


def download_audio(url: str, output_dir: str) -> str:
    if APIFY_TOKEN and is_youtube_url(url):
        try:
            return download_audio_via_apify(url, output_dir)
        except Exception as exc:
            raise RuntimeError(f"Apify YouTube download failed: {exc}") from exc

    yt_dlp = get_yt_dlp_module()
    ensure_ffmpeg()

    output_dir_path = Path(output_dir)
    output_dir_path.mkdir(parents=True, exist_ok=True)
    cookies_file = maybe_write_youtube_cookies_file(output_dir_path)
    output_template = str(output_dir_path / "downloaded.%(ext)s")
    ydl_base_opts = {
        "outtmpl": output_template,
        "quiet": True,
        "no_warnings": True,
        "noprogress": True,
        "socket_timeout": 30,
        "retries": 3,
        "fragment_retries": 3,
        "file_access_retries": 3,
        "extractor_retries": 3,
        "geo_bypass": True,
        "logger": QuietYtdlpLogger(),
        "http_headers": {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/127.0.0.0 Safari/537.36"
            )
        },
        "extractor_args": {
            "youtube": {
                "player_client": ["web_embedded", "web_creator", "web"]
            }
        },
        "postprocessors": [
            {
                "key": "FFmpegExtractAudio",
                "preferredcodec": "wav",
                "preferredquality": "192",
            }
        ],
    }
    if cookies_file:
        ydl_base_opts["cookiefile"] = str(cookies_file)
    proxy_url = get_proxy_url("https")
    if proxy_url:
        ydl_base_opts["proxy"] = proxy_url
    try:
        with yt_dlp.YoutubeDL({**ydl_base_opts, "skip_download": True}) as ydl:
            info = ydl.extract_info(url, download=False)
        selected_format_id = select_best_audio_format_id(info)
        ydl_opts = dict(ydl_base_opts)
        ydl_opts["format"] = selected_format_id or "bestaudio/best"

        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            ydl.download([url])
    except yt_dlp.utils.DownloadError as exc:
        error_text = str(exc)
        if "Video unavailable" in error_text:
            raise RuntimeError("This YouTube video is unavailable or private.") from exc
        if "Sign in to confirm your age" in error_text:
            raise RuntimeError("This YouTube video is age-restricted and cannot be downloaded without cookies.") from exc
        if "Please sign in" in error_text or "cookies" in error_text.lower():
            raise RuntimeError(
                "This YouTube video requires authenticated cookies. Add a Railway env var named "
                "YOUTUBE_COOKIES_B64 containing base64-encoded Netscape YouTube cookies, then redeploy."
            ) from exc
        if "Requested format is not available" in error_text:
            raise RuntimeError(
                "The video is available, but YouTube did not expose a compatible downloadable audio format to the server. "
                "Please retry once, try another video, or upload the media file directly."
            ) from exc
        if "HTTP Error 429" in error_text:
            raise RuntimeError("YouTube is rate-limiting the server right now. Please retry in a moment or upload the file directly.") from exc
        if "HTTP Error 403" in error_text or "PO Token" in error_text or "challenge" in error_text:
            raise RuntimeError(
                "YouTube blocked this server request. The app has been updated with stronger download support; please retry once. "
                "If the video still fails, upload the media file directly."
            ) from exc
        raise RuntimeError(
            f"Could not download audio from that YouTube URL. Details: {error_text}"
        ) from exc

    wav_files = sorted(output_dir_path.glob("downloaded*.wav"))
    if not wav_files:
        raise RuntimeError("Audio download completed, but no WAV file was produced.")
    return str(wav_files[0])


def extract_audio_from_video(video_path: str, output_audio_path: str) -> str:
    video_file_clip = get_video_file_clip_class()
    ensure_ffmpeg()

    with video_file_clip(video_path) as clip:
        if clip.audio is None:
            raise RuntimeError("The uploaded video does not contain an audio track.")
        clip.audio.write_audiofile(output_audio_path, logger=None)
    return output_audio_path


def transcribe_audio(audio_path: str) -> str:
    model = get_whisper_model()
    result = model.transcribe(audio_path, fp16=False, verbose=False)
    return result.get("text", "")


def build_transcript_cache_key(source_identifier: str) -> str:
    digest = hashlib.sha1(source_identifier.encode("utf-8")).hexdigest()
    return digest


def load_cached_transcript(cache_key: str) -> Optional[str]:
    cache_path = TRANSCRIPT_CACHE_DIR / f"{cache_key}.txt"
    if not cache_path.exists():
        return None
    try:
        text = cache_path.read_text(encoding="utf-8").strip()
        return text or None
    except Exception:
        return None


def save_cached_transcript(cache_key: str, transcript: str) -> None:
    try:
        TRANSCRIPT_CACHE_DIR.mkdir(parents=True, exist_ok=True)
        (TRANSCRIPT_CACHE_DIR / f"{cache_key}.txt").write_text(transcript, encoding="utf-8")
    except Exception:
        return None


def split_audio_for_transcription(audio_path: str, output_dir: str, chunk_seconds: int = TRANSCRIPTION_CHUNK_SECONDS) -> list[str]:
    ensure_ffmpeg()
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    segment_pattern = str(output_path / "segment_%03d.wav")
    command = [
        "ffmpeg",
        "-y",
        "-i",
        audio_path,
        "-vn",
        "-ac",
        "1",
        "-ar",
        "16000",
        "-f",
        "segment",
        "-segment_time",
        str(chunk_seconds),
        segment_pattern,
    ]
    completed = subprocess.run(command, capture_output=True, text=True)
    if completed.returncode != 0:
        return [audio_path]
    chunks = sorted(str(path) for path in output_path.glob("segment_*.wav"))
    return chunks or [audio_path]


def transcribe_audio_with_fallback(
    audio_path: str,
    *,
    cache_key: Optional[str] = None,
    progress_callback=None,
) -> str:
    profiler = StageProfiler(f"transcription:{Path(audio_path).name}")
    if cache_key:
        with profiler.stage("transcript_cache_lookup"):
            cached = load_cached_transcript(cache_key)
        if cached:
            if progress_callback:
                progress_callback("transcribing_audio", "Using cached transcript.", 45)
            profiler.log_total()
            return cached

    with tempfile.TemporaryDirectory(prefix="transcription-chunks-") as temp_dir:
        with profiler.stage("audio_chunking"):
            chunks = split_audio_for_transcription(audio_path, temp_dir)
        transcript_parts: list[str] = []
        total = len(chunks)
        for index, chunk_path in enumerate(chunks, start=1):
            if progress_callback:
                progress = min(70, 20 + int((index / max(total, 1)) * 45))
                progress_callback(
                    "transcribing_audio",
                    f"Transcribing audio part {index} of {total}...",
                    progress,
                )
            last_error: Optional[Exception] = None
            for _ in range(2):
                try:
                    with profiler.stage(f"transcribe_chunk_{index}"):
                        chunk_text = transcribe_audio(chunk_path).strip()
                    if chunk_text:
                        transcript_parts.append(chunk_text)
                    last_error = None
                    break
                except Exception as exc:
                    last_error = exc
            if last_error is not None:
                raise RuntimeError(f"Transcription failed for audio part {index} of {total}.") from last_error

    with profiler.stage("transcript_cleaning"):
        transcript = clean_source_text(" ".join(transcript_parts))
    if not transcript.strip():
        raise RuntimeError("Transcription failed to produce readable text.")
    if cache_key:
        with profiler.stage("transcript_cache_store"):
            save_cached_transcript(cache_key, transcript)
    profiler.log_total()
    return transcript


def chunk_text(text: str, max_chars: int = 600) -> list[str]:
    try:
        sentences = sent_tokenize(text)
    except LookupError:
        # Fall back when punkt data is unavailable in serverless environments.
        sentences = [segment.strip() for segment in re.split(r"(?<=[.!?])\s+", text) if segment.strip()]
    chunks, current = [], ""
    for sentence in sentences:
        if len(current) + len(sentence) <= max_chars:
            current += " " + sentence
        else:
            if current.strip():
                chunks.append(current.strip())
            current = sentence
    if current.strip():
        chunks.append(current.strip())
    return chunks or [text[:max_chars]]


def get_embeddings(chunks: list[str]) -> np.ndarray:
    embedder = get_embedder()
    return embedder.encode(chunks, convert_to_numpy=True, show_progress_bar=False)


def build_faiss_index(embeddings: np.ndarray):
    faiss = get_faiss_module()
    dim = len(embeddings[0])
    index = faiss.IndexFlatL2(dim)
    index.add(np.array(embeddings).astype("float32"))
    return index


def retrieve_context(query: str, chunks: list[str], index, k: int = 3) -> str:
    embedder = get_embedder()
    query_embedding = embedder.encode([query], convert_to_numpy=True)
    _, indices = index.search(query_embedding.astype("float32"), k)
    return "\n".join([chunks[i] for i in indices[0]])


def split_sentences(text: str) -> list[str]:
    normalized = " ".join((text or "").split())
    if not normalized:
        return []
    try:
        return [sentence.strip() for sentence in sent_tokenize(normalized) if sentence.strip()]
    except LookupError:
        return [sentence.strip() for sentence in re.split(r"(?<=[.!?])\s+", normalized) if sentence.strip()]


def extract_top_keywords(text: str, limit: int = 5) -> list[str]:
    tokens = re.findall(r"[A-Za-z][A-Za-z0-9'-]{3,}", text or "")
    counts = Counter(token.lower() for token in tokens if token.lower() not in COMMON_TOPIC_STOPWORDS)
    keywords: list[str] = []
    for word, _ in counts.most_common(limit * 3):
        cleaned = word.strip("-'")
        if not cleaned or cleaned in keywords:
            continue
        keywords.append(cleaned.title())
        if len(keywords) >= limit:
            break
    return keywords


def build_local_headline(query: str, transcript: str) -> str:
    points = build_local_summary_points(transcript, max_points=3)
    if len(points) > 1 and len(points[0].split()) < 10:
        points = points[1:]
    keywords = extract_top_keywords(" ".join(points), limit=3)
    if keywords:
        return f"{' | '.join(keywords)}: {query}".replace("Give ", "").strip()
    return f"Key Takeaways: {query}".strip()


def build_local_summary_points(transcript: str, max_points: int = 4) -> list[str]:
    sentences = split_sentences(transcript)
    if not sentences:
        return ["No transcript could be extracted from the provided source."]

    selected: list[str] = []
    for sentence in sentences:
        compact = " ".join(sentence.split())
        if len(compact) < 40:
            continue
        if compact in selected:
            continue
        selected.append(compact)
        if len(selected) >= max_points:
            break

    if not selected:
        selected = sentences[:max_points]
    return selected[:max_points]


def build_local_topic_points(topic_title: str, support_text: str, importance_seed: str) -> list[dict[str, str]]:
    seed_sentences = split_sentences(f"{support_text} {importance_seed}") or [support_text or importance_seed]
    labels = [
        "Core argument",
        "Why it matters",
        "Practical relevance",
    ]
    points: list[dict[str, str]] = []
    for index, sentence in enumerate(seed_sentences[:3]):
        cleaned_sentence = remove_timestamps(sanitize_inline_html(sentence[:260].strip()))
        if not cleaned_sentence:
            continue
        points.append(
            {
                "label": labels[index] if index < len(labels) else f"{topic_title} insight",
                "description": cleaned_sentence,
            }
        )
    return points


def build_local_topic_details(summary_text: str) -> list[dict[str, object]]:
    points = [line[2:].strip() for line in summary_text.splitlines() if line.startswith("- ")]
    if len(points) > 1 and len(points[0].split()) < 10:
        points = points[1:]
    topic_source = " ".join(points) if points else summary_text
    keywords = extract_top_keywords(topic_source, limit=8)
    details: list[dict[str, str]] = []
    support_pool = points or split_sentences(summary_text)
    for index, keyword in enumerate(keywords[:8]):
        support_text = support_pool[min(index, len(support_pool) - 1)] if support_pool else summary_text
        title = build_editorial_topic_title(keyword, support_text, index)
        importance_seed = support_pool[(index + 1) % len(support_pool)] if support_pool else support_text
        details.append(
            {
                "title": title,
                "explanation": remove_timestamps(sanitize_inline_html(
                    f"{support_text[:220].strip()} This angle gives the writer enough substance to build a full article instead of a thin summary."
                )),
                "importance": remove_timestamps(sanitize_inline_html(
                    f"It matters because it opens a clearer view of the argument, the stakes, and the examples that give {keyword.lower()} real weight. {importance_seed[:140].strip()}"
                )),
                "points": build_local_topic_points(title, support_text, importance_seed),
            }
        )
    return details or [{
        "id": "topic_1",
        "title": "The Main Argument Behind the Story",
        "summary": sanitize_inline_html("The source centers on one clearly expandable issue with enough direction, detail, and tension to support a professional article."),
        "explanation": sanitize_inline_html("The source centers on one clearly expandable issue with enough direction, detail, and tension to support a professional article."),
        "importance": sanitize_inline_html("It matters because it captures the strongest idea in the material and gives the writer a solid editorial angle to build on."),
        "points": [
            {
                "label": "Main lesson",
                "description": "The source presents one central idea strongly enough to support a focused, publication-ready article.",
            }
        ],
    }]


def fallback_summary(transcript: str, query: str) -> dict[str, str]:
    points = build_local_summary_points(transcript)
    return {
        "headline": build_local_headline(query, transcript),
        "summary": "\n".join(f"- {point}" for point in points),
    }


def fallback_topics(summary_text: str) -> list[str]:
    points = [line[2:].strip() for line in summary_text.splitlines() if line.startswith("- ")]
    if len(points) > 1 and len(points[0].split()) < 10:
        points = points[1:]
    topic_source = " ".join(points) if points else summary_text
    return extract_top_keywords(topic_source, limit=5) or ["Main Topic"]


def get_topic_details_from_summary(summary_text: str) -> list[dict[str, object]]:
    prompt = f"""
You are an expert video analyst and editorial researcher.

Analyze the summary below and create 5 to 8 meaningful topic options based strictly on the actual source.
Each topic must include:
- title
- explanation
- importance
- points: 2 to 3 short sub-points with:
  - label
  - description

Rules:
- The title must sound like a real analytical angle, not a one-word label.
- Avoid generic titles such as "News", "Update", "Main Topic", or "Overview".
- Keep titles specific, useful, and directly connected to the summary.
- Explanation must briefly say what the topic means in the source.
- Importance must explain why the angle matters.
- Each point should feel like deep analysis, not a transcript line.
- Do not include timestamps.
- Do not repeat the same idea in different wording.

Return valid JSON only in this shape:
[
  {{
    "title": "Topic title",
    "explanation": "Short explanation",
    "importance": "Why it matters",
    "points": [
      {{"label": "Sub-point heading", "description": "Short explanation based on the source"}},
      {{"label": "Sub-point heading", "description": "Short explanation based on the source"}}
    ]
  }}
]

Summary:
{summary_text}
"""
    try:
        raw_text = gemini_generate_text(prompt)
        payload = json.loads(raw_text)
        if not isinstance(payload, list):
            raise ValueError("Topic details payload must be a list.")
        details: list[dict[str, object]] = []
        for item in payload:
            if not isinstance(item, dict):
                continue
            title = str(item.get("title", "")).strip()
            explanation = str(item.get("explanation", "")).strip()
            importance = str(item.get("importance", "")).strip()
            if not title and not explanation:
                continue
            cleaned_title = clean_topic_title(title)
            if title_needs_editorial_rewrite(cleaned_title):
                cleaned_title = build_editorial_topic_title(cleaned_title or "Main Topic", explanation or importance, len(details))
            points = []
            for raw_point in item.get("points", []) if isinstance(item.get("points", []), list) else []:
                cleaned_point = sanitize_topic_point(raw_point)
                if cleaned_point:
                    points.append(cleaned_point)
            details.append(
                {
                    "id": f"topic_{len(details) + 1}",
                    "title": cleaned_title,
                    "summary": remove_timestamps(sanitize_inline_html(explanation)),
                    "explanation": remove_timestamps(sanitize_inline_html(explanation)),
                    "importance": remove_timestamps(sanitize_inline_html(importance)),
                    "points": points,
                }
            )
        return details[:8] or build_local_topic_details(summary_text)
    except Exception:
        return build_local_topic_details(summary_text)


def fallback_article(headline_text: str, summary_text: str, topic: Optional[str] = None) -> str:
    topic_title = topic or "Main Topic"
    summary_points = [line[2:].strip() for line in summary_text.splitlines() if line.startswith("- ")]
    body_points = summary_points or ["No summary details were available from the source."]
    lead = body_points[0]
    support = body_points[1] if len(body_points) > 1 else lead
    detail = body_points[2] if len(body_points) > 2 else support
    closing_support = body_points[3] if len(body_points) > 3 else detail
    section_one, section_two, section_three, section_four = build_dynamic_subheadings(topic_title, body_points, "Blog Article")
    return postprocess_article_html(
        f"""
<h2>{headline_text}</h2>
<p>{lead} That opening detail immediately gives <strong>{topic_title}</strong> a sense of consequence. It suggests a story that is not only timely, but layered enough to reward closer attention. The point is not simply that something happened. The point is that the development carries implications, and those implications become clearer once the surrounding examples are placed side by side.</p>
<p>{support} Read together, the early signals create the outline of a bigger argument. They point toward a shift in priorities, a contest of ideas, or a change in public mood that deserves more than a passing mention. That is where the article gains its strength: it treats the subject as something substantial enough to explain, not merely something dramatic enough to repeat.</p>
<h3>{section_one}</h3>
<p>{detail} One of the most persuasive qualities of <strong>{topic_title}</strong> is the way the evidence accumulates. Each example adds pressure to the same central claim, making the story feel grounded rather than speculative. When readers can see how the argument is supported step by step, the writing becomes more convincing without having to overstate its point.</p>
<p>{support} That structure matters. Strong editorial writing does not rely on noise. It relies on sequence, emphasis, and proportion. Here, the sequence gives the subject momentum, while the emphasis makes clear which ideas deserve the reader's attention first.</p>
<h3>{section_two}</h3>
<p>{lead} {support} These details work because they turn an abstract theme into something tangible. They give the reader enough specificity to understand what is at stake, who is affected, and why the issue feels larger than a single moment. That is the difference between a thin recap and a serious article.</p>
<p>{detail} {closing_support} Once these details are allowed to sit together, the topic begins to reveal its broader shape. It becomes easier to see the direction of the argument, the pressures behind it, and the reasons it may continue to matter long after the immediate event fades from view.</p>
<h3>{section_three}</h3>
<p><strong>{topic_title}</strong> carries weight because it opens into a larger conversation. It touches the logic behind decisions, the values guiding those decisions, and the public consequences that follow. Topics with that range rarely stay confined to one headline. They tend to echo across policy, culture, business, or civic life, depending on the material behind them.</p>
<p>That is also what makes the article readable. Readers are not being handed disconnected talking points. They are being given a coherent interpretation of why the issue deserves sustained attention. The material holds together because the examples reinforce the same core direction instead of competing with one another.</p>
<h3>{section_four}</h3>
<p>{closing_support} By the end, the subject feels more fully formed. The strongest impression is not simply that the topic is relevant, but that it offers a useful way to understand a bigger pattern already taking shape. That makes it a stronger editorial subject than a one-dimensional trend piece or a narrow update.</p>
<p>The article lands best when it leaves the reader with clarity rather than volume. That clarity comes from connecting examples to meaning, and from treating <strong>{topic_title}</strong> as a live issue with lasting significance. When the writing does that well, the subject no longer feels like raw source material. It feels like a finished argument.</p>
"""
    )


def fallback_article_for_type(
    headline_text: str,
    summary_text: str,
    topic: Optional[str],
    article_type: str,
) -> str:
    base = fallback_article(headline_text, summary_text, topic)
    if article_type == "News Article":
        topic_title = topic or "the story"
        points = [line[2:].strip() for line in summary_text.splitlines() if line.startswith("- ")] or [summary_text]
        lead = points[0]
        context = points[1] if len(points) > 1 else lead
        update = points[2] if len(points) > 2 else context
        section_one, section_two, _, section_four = build_dynamic_subheadings(topic_title, points, "News Article")
        return postprocess_article_html(
            f"""
<h2>{headline_text}</h2>
<p>{lead} The development places <strong>{topic_title}</strong> at the center of the story and frames the issue as one with immediate relevance beyond a single update.</p>
<h3>{section_one}</h3>
<p>{context} The core facts indicate a pattern that is still unfolding, which is why the subject warrants direct, sober reporting rather than promotional framing.</p>
<h3>{section_two}</h3>
<p>{update} Taken together, the available details provide the background needed to understand why this development matters now and how it connects to the broader context surrounding it.</p>
<h3>{section_four}</h3>
<p>The clearest takeaway is that the story should be followed for what it reveals over time, not just for its immediate headline value. That makes it a stronger news subject than a one-off talking point.</p>
"""
        )
    return base


def proofread_article_html(
    article_html: str,
    *,
    article_type: str,
    topic: str,
) -> str:
    prompt = f"""
You are proofreading an article before publication.

Article type:
{article_type}

Topic:
{topic}

Instructions:
- Improve grammar, punctuation, and flow.
- Keep the facts, meaning, and structure intact.
- Remove repeated sentences, filler openings, and awkward transitions.
- Do not add new facts.
- Return clean HTML only using <h2>, <h3>, <p>, <ul>, <li>, <strong>.

Article HTML:
{article_html}
"""
    try:
        return postprocess_article_html(gemini_generate_text(prompt)) or postprocess_article_html(article_html)
    except Exception:
        return postprocess_article_html(article_html)


def article_needs_refinement(article_html: str) -> bool:
    text = strip_html_tags(article_html)
    if len(text.split()) < 420:
        return True
    headings = re.findall(r"<h[23]>", article_html)
    if len(headings) < 3:
        return True
    sentences = split_sentences(text)
    if not sentences:
        return False
    unique_ratio = len(set(sentence.lower() for sentence in sentences)) / max(len(sentences), 1)
    return unique_ratio < 0.75


def generate_article_html(
    *,
    headline_text: str,
    summary_text: str,
    topic: Optional[str],
    topic_summary: str = "",
    topic_points: list[dict[str, str]] | None = None,
    article_type: str,
    source_context: str,
    target_audience: str,
    variant_index: int = 0,
    profiler: Optional[StageProfiler] = None,
) -> str:
    article_type = article_type if article_type in ARTICLE_TYPE_INSTRUCTIONS else "Blog Article"
    topic_text = topic or "Main Topic"
    source_excerpt = clean_source_text(source_context or summary_text)[:2800]
    normalized_topic_points = topic_points or []
    topic_points_text = "\n".join(
        f"- {point.get('label', '').strip()}: {point.get('description', '').strip()}"
        for point in normalized_topic_points
        if isinstance(point, dict) and (point.get("label") or point.get("description"))
    )
    keywords = extract_top_keywords(f"{summary_text} {topic_text} {topic_summary} {topic_points_text} {source_excerpt}", limit=8)
    uniqueness_instruction = (
        "Use a fresh structure and angle for this version."
        if variant_index == 0
        else f"Create a distinct variation #{variant_index + 1} with a noticeably different angle, examples emphasis, and subheading sequence."
    )
    prompt = f"""
You are a senior editorial writer and AI content strategist.

Selected article type:
{article_type}

Topic:
{topic_text}

Selected topic summary:
{topic_summary or "Use the source context to infer the strongest angle for this topic."}

Selected topic sub-points:
{topic_points_text or "- Use the strongest source-backed details relevant to the selected topic."}

Target audience:
{target_audience}

Source summary:
{summary_text}

Cleaned source context:
{source_excerpt}

Supporting keywords:
{", ".join(keywords)}

Instructions:
- {ARTICLE_TYPE_INSTRUCTIONS[article_type]}
- Base the article strictly on the source context and summary above.
- Do not copy transcript phrasing directly.
- Remove filler wording, transcript noise, timestamps, speaker labels, and repetition.
- Use the selected topic as the center of the article.
- Generate unique headings and subheadings based on the actual source content.
- Keep the writing polished, publication-ready, and natural.
- Avoid generic filler openings such as "In today's digital world" or "It is important to note."
- {uniqueness_instruction}
- Return clean HTML only using <h2>, <h3>, <p>, <ul>, <li>, <strong>.
"""
    try:
        if profiler:
            with profiler.stage("article_generation"):
                article_text = gemini_generate_text(prompt)
        else:
            article_text = gemini_generate_text(prompt)
        article_html = postprocess_article_html(article_text) or fallback_article_for_type(headline_text, summary_text, topic_text, article_type)
    except Exception:
        article_html = fallback_article_for_type(headline_text, summary_text, topic_text, article_type)

    if profiler:
        with profiler.stage("duplicate_cleanup"):
            article_html = deduplicate_article_html(article_html)
    else:
        article_html = deduplicate_article_html(article_html)

    if ENABLE_DEEP_ARTICLE_REFINEMENT and article_needs_refinement(article_html):
        if profiler:
            with profiler.stage("proofreading_refinement"):
                return proofread_article_html(article_html, article_type=article_type, topic=topic_text)
        return proofread_article_html(article_html, article_type=article_type, topic=topic_text)
    return article_html


def build_fast_context_from_transcript(transcript: str, chunks: list[str]) -> str:
    if transcript.strip():
        normalized = " ".join(transcript.split())
        return normalized[:FAST_ANALYSIS_TRANSCRIPT_LIMIT]
    return " ".join(chunks[:FAST_ANALYSIS_CHUNK_LIMIT])


def build_article_image_url(topic: Optional[str]) -> str:
    topic_text = (topic or "news article").lower()
    image_map = [
        ({"technology", "ai", "software", "digital", "cyber"}, "https://images.unsplash.com/photo-1518770660439-4636190af475?auto=format&fit=crop&w=1200&q=80"),
        ({"finance", "business", "market", "economy", "trade"}, "https://images.unsplash.com/photo-1520607162513-77705c0f0d4a?auto=format&fit=crop&w=1200&q=80"),
        ({"media", "video", "film", "broadcast", "news"}, "https://images.unsplash.com/photo-1495020689067-958852a7765e?auto=format&fit=crop&w=1200&q=80"),
        ({"health", "medical", "hospital", "wellness"}, "https://images.unsplash.com/photo-1505751172876-fa1923c5c528?auto=format&fit=crop&w=1200&q=80"),
        ({"sports", "football", "cricket", "game", "team"}, "https://images.unsplash.com/photo-1517649763962-0c623066013b?auto=format&fit=crop&w=1200&q=80"),
        ({"travel", "tourism", "flight", "hotel"}, "https://images.unsplash.com/photo-1507525428034-b723cf961d3e?auto=format&fit=crop&w=1200&q=80"),
    ]
    for keywords, url in image_map:
        if any(keyword in topic_text for keyword in keywords):
            return url
    return "https://images.unsplash.com/photo-1504711434969-e33886168f5c?auto=format&fit=crop&w=1200&q=80"


def gemini_rag(context: str, query: str, *, source_url: str = "", retry_strict: bool = False) -> dict[str, object]:
    retry_warning = (
        "The previous response summarized metadata instead of the video content. Re-analyze the actual video content only. Ignore URLs, credits, descriptions, and metadata.\n"
        if retry_strict else ""
    )
    source_reference = f"\nSource URL:\n{source_url}\n" if source_url else ""
    prompt = f"""
You are an expert video analyst and editorial researcher.

{retry_warning}
Analyze the actual content of the provided source deeply. Do not transcribe it. Identify the central idea, summarize it clearly, and break it into meaningful topics.

Context from media:
{context}
{source_reference}

User Query:
{query}

Instructions:
- Return only valid JSON.
- Create one strong, specific heading for the source.
- Write one concise but meaningful summary paragraph.
- Create 3 to 7 key points from the actual source.
- Create 4 to 7 dynamic topics based on the actual source.
- Each topic must include 2 to 3 important sub-points with short explanations.
- Keep the wording natural, professional, analytical, and grounded in the source.
- Do not invent facts that are not supported by the source.
- Do not include timestamps.
- Do not summarize the YouTube description, links, credits, tags, channel metadata, Instagram URLs, promotional text, or repeated keywords.
- Ignore all URLs and boilerplate metadata unless the actual spoken content is directly about them.
- Do not sound like a transcript.
- Do not repeat the same point.
- Use the user's instruction to decide the angle of the analysis.

Return JSON in this shape:
{{
  "heading": "Main heading",
  "summary": "Short meaningful summary",
  "key_points": ["Point one", "Point two"],
  "topics": [
    {{
      "id": "topic_1",
      "title": "Topic title",
      "summary": "Short topic summary",
      "importance": "Why this topic matters",
      "points": [
        {{"label": "Sub-point heading", "description": "Short explanation based on the source"}}
      ]
    }}
  ]
}}
"""
    response_text = gemini_generate_text(prompt)
    payload = json.loads(response_text)
    if not isinstance(payload, dict):
        raise ValueError("Gemini analysis payload must be an object.")

    heading = remove_timestamps(sanitize_inline_html(str(payload.get("heading", "")).strip())) or "Media summary"
    summary = remove_timestamps(sanitize_inline_html(str(payload.get("summary", "")).strip()))
    raw_key_points = payload.get("key_points", [])
    key_points = []
    if isinstance(raw_key_points, list):
        for item in raw_key_points:
            cleaned = remove_timestamps(sanitize_inline_html(str(item).strip()))
            if cleaned and cleaned not in key_points:
                key_points.append(cleaned)

    raw_topics = payload.get("topics", [])
    topics: list[dict[str, object]] = []
    if isinstance(raw_topics, list):
        for item in raw_topics:
            if not isinstance(item, dict):
                continue
            title = clean_topic_title(str(item.get("title", "")).strip())
            explanation = remove_timestamps(sanitize_inline_html(str(item.get("summary", "") or item.get("explanation", "")).strip()))
            importance = remove_timestamps(sanitize_inline_html(str(item.get("importance", "")).strip()))
            points = []
            for raw_point in item.get("points", []) if isinstance(item.get("points", []), list) else []:
                cleaned_point = sanitize_topic_point(raw_point)
                if cleaned_point:
                    points.append(cleaned_point)
            if not title:
                continue
            topics.append(
                {
                    "id": str(item.get("id", "")).strip() or f"topic_{len(topics) + 1}",
                    "title": title,
                    "summary": explanation,
                    "explanation": explanation,
                    "importance": importance,
                    "points": points,
                }
            )

    if not summary and key_points:
        summary = " ".join(key_points[:2])

    result = {
        "heading": heading,
        "headline": heading,
        "summary": summary or "No summary generated.",
        "key_points": key_points,
        "topics": topics,
    }
    if not analysis_output_is_valid(result):
        if retry_strict:
            raise ValueError("Gemini analysis output was polluted by metadata.")
        return gemini_rag(context, query, source_url=source_url, retry_strict=True)
    return result


def direct_gemini_youtube_analysis(url: str, query: str, metadata: Optional[dict[str, object]] = None) -> dict[str, object]:
    model_name = get_gemini_model_name()
    title_hint = sanitize_youtube_metadata_text(str((metadata or {}).get("title") or ""))
    context = (
        f"Video title hint: {title_hint}\n"
        "Analyze the actual YouTube video content from the provided source URL.\n"
        "Ignore metadata, description, URLs, credits, tags, channel info, and social media links.\n"
    )
    logger.info(
        "youtube-direct-analysis-attempt | model=%s | request_type=text+url | url=%s",
        model_name,
        url,
    )
    try:
        result = gemini_rag(context, query, source_url=url)
        logger.info(
            "youtube-direct-analysis-success | model=%s | request_type=text+url | url=%s",
            model_name,
            url,
        )
        return result
    except Exception as exc:
        logger.warning(
            "youtube-direct-analysis-failed | model=%s | request_type=text+url | url=%s | error_type=%s | error=%s",
            model_name,
            url,
            exc.__class__.__name__,
            str(exc),
        )
        raise RuntimeError(direct_gemini_failure_reason(exc)) from exc


def generate_news_article(headline_text: str, summary_text: str, topic: Optional[str] = None) -> str:
    if not headline_text.strip():
        headline_text = "Media summary"
    return generate_article_html(
        headline_text=headline_text,
        summary_text=summary_text,
        topic=topic,
        article_type="News Article",
        source_context=summary_text,
        target_audience="General readers",
        variant_index=0,
    )


def get_topics_from_summary(summary_text: str) -> list[str]:
    return [item["title"] for item in get_topic_details_from_summary(summary_text)]


def summary_to_key_points(summary_text: str) -> list[str]:
    return [line[2:].strip() for line in summary_text.splitlines() if line.strip().startswith("- ")]


def enrich_analysis(
    result: dict[str, str],
    generate_article: bool = False,
    selected_topics: Optional[list[str]] = None,
    article_count: int = 1,
    article_type: str = "Blog Article",
    target_audience: str = "General readers",
    source_context: str = "",
    source_type: str = "url",
    source_cache_key: str = "",
) -> dict[str, object]:
    if source_cache_key and len(source_context or "") < 1200:
        cached_source = load_cached_source_content(source_cache_key)
        if isinstance(cached_source, dict) and isinstance(cached_source.get("content"), str):
            source_context = str(cached_source["content"])
    headline = str(result.get("heading") or result.get("headline") or "Media summary")
    summary = remove_timestamps(str(result.get("summary", "") or ""))
    raw_key_points = result.get("key_points", [])
    if isinstance(raw_key_points, list):
        key_points = [remove_timestamps(sanitize_inline_html(str(item))) for item in raw_key_points if str(item).strip()]
    else:
        key_points = []
    if not key_points:
        key_points = summary_to_key_points(summary)
    structured_topics = result.get("topics", []) if isinstance(result.get("topics", []), list) else None
    topic_details, used_fallback_topics = build_topic_details_bundle(summary, structured_topics=structured_topics)
    topics = [item["title"] for item in topic_details]
    payload: dict[str, object] = {
        "heading": headline,
        "headline": headline,
        "summary": summary,
        "key_points": key_points,
        "topics": topics,
        "topic_details": topic_details,
        "articles": [],
        "article_type": article_type,
        "target_audience": target_audience,
        "source_type": source_type,
        "source_context_preview": clean_source_text(source_context or summary)[:700],
        "source_cache_key": source_cache_key,
        "topic_generation_warning": (
            "Topic ideas were simplified because the source content could not be fully expanded into richer angles."
            if used_fallback_topics else ""
        ),
    }

    if not generate_article:
        return payload

    topics_to_use = (selected_topics[:1] if selected_topics else topics[:1]) or ["Main topic"]
    topic_detail_lookup = {
        str(item.get("title", "")).strip().lower(): item
        for item in topic_details
        if isinstance(item, dict) and str(item.get("title", "")).strip()
    }
    work_items: list[tuple[int, str, int]] = []
    for topic_index, topic in enumerate(topics_to_use):
        for variant_index in range(max(article_count, 1)):
            work_items.append(((topic_index * max(article_count, 1)) + variant_index, topic, variant_index))

    def build_one(item: tuple[int, str, int]) -> tuple[int, dict[str, object]]:
        order, topic_name, variant_index = item
        topic_detail = topic_detail_lookup.get(topic_name.strip().lower(), {})
        return order, build_article_package(
            headline_text=headline,
            summary_text=summary,
            topic=topic_name,
            topic_summary=str(topic_detail.get("summary") or topic_detail.get("explanation") or ""),
            topic_points=topic_detail.get("points") if isinstance(topic_detail.get("points"), list) else [],
            article_type=article_type,
            source_context=source_context or summary,
            target_audience=target_audience,
            variant_index=order,
        )

    articles: list[Optional[dict[str, object]]] = [None] * len(work_items)
    max_workers = min(2, len(work_items)) if work_items else 1
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        for order, payload_item in executor.map(build_one, work_items):
            articles[order] = payload_item
    payload["articles"] = [item for item in articles if item is not None]
    return payload


def build_analysis_context(transcript: str, chunks: list[str], query: str) -> str:
    if FAST_ANALYSIS_MODE:
        return build_fast_context_from_transcript(transcript, chunks)

    embeddings = get_embeddings(chunks)
    index = build_faiss_index(embeddings)
    return retrieve_context(query, chunks, index)


def summarize_transcript_chunk(chunk: str, query: str) -> str:
    cache_key = build_summary_cache_key(chunk, query)
    cached = load_cached_summary(cache_key)
    if cached:
        logger.info("chunk-summary-cache-hit | %.0f chars", len(chunk))
        return cached

    prompt = f"""
You are preparing source notes for article generation.

User request:
{query}

Transcript chunk:
{chunk}

Return 3 to 5 concise bullet points that capture the factual ideas, arguments, examples, and entities in this chunk.
Keep the wording clean and remove filler or transcript noise.
"""
    try:
        raw = gemini_generate_text(prompt)
        lines = [line.strip() for line in raw.splitlines() if line.strip()]
        bullets = [line if line.startswith("- ") else f"- {line.lstrip('- ').strip()}" for line in lines[:5]]
        summary_text = "\n".join(bullets)
    except Exception:
        summary_text = "\n".join(f"- {point}" for point in build_local_summary_points(chunk, max_points=4))
    save_cached_summary(cache_key, summary_text)
    return summary_text


def build_chunk_safe_analysis_input(transcript: str, query: str, progress_callback=None) -> str:
    transcript = clean_source_text(transcript)
    if len(transcript) <= LONG_TRANSCRIPT_THRESHOLD:
        return transcript

    transcript_chunks = chunk_text(transcript, max_chars=2800)
    total = len(transcript_chunks)
    summaries: list[str] = [""] * total

    def summarize_one(args: tuple[int, str]) -> tuple[int, str]:
        index, chunk = args
        return index, summarize_transcript_chunk(chunk, query)

    work_items = list(enumerate(transcript_chunks))
    if len(work_items) == 1:
        index, summary = summarize_one(work_items[0])
        summaries[index] = summary
        if progress_callback:
            progress_callback("chunk_summarization", "Summarizing transcript chunk 1 of 1...", 82)
    else:
        with ThreadPoolExecutor(max_workers=min(SUMMARY_CONCURRENCY, len(work_items))) as executor:
            completed = 0
            for index, summary in executor.map(summarize_one, work_items):
                summaries[index] = summary
                completed += 1
                if progress_callback:
                    progress = min(88, 72 + int((completed / max(total, 1)) * 14))
                    progress_callback(
                        "chunk_summarization",
                        f"Summarizing transcript chunk {completed} of {total}...",
                        progress,
                    )
    return "\n".join(summaries)


def analyze_media(
    video_path: Optional[str] = None,
    local_audio_path: Optional[str] = None,
    query: str = "Summarize the content",
    *,
    transcript_cache_key: Optional[str] = None,
    progress_callback=None,
) -> tuple[dict[str, str], str]:
    profiler = StageProfiler(f"analyze-media:{Path(video_path or local_audio_path or 'source').name}")
    audio_source_path = None
    if local_audio_path:
        audio_source_path = local_audio_path
    elif video_path:
        with tempfile.TemporaryDirectory(prefix="media-analyzer-extract-") as temp_dir:
            extracted_audio_path = str(Path(temp_dir) / "extracted_audio.wav")
            with profiler.stage("source_fetch_download"):
                audio_source_path = extract_audio_from_video(video_path, extracted_audio_path)
            if progress_callback:
                progress_callback("transcribing_audio", "Transcribing extracted audio...", 35)
            try:
                with profiler.stage("transcription"):
                    transcript = transcribe_audio_with_fallback(audio_source_path, cache_key=transcript_cache_key, progress_callback=progress_callback)
            except RuntimeError as exc:
                if "readable text" in str(exc).lower():
                    profiler.log_total()
                    return fallback_summary("", query), ""
                raise
            if not transcript.strip():
                profiler.log_total()
                return fallback_summary("", query), ""
            if progress_callback:
                progress_callback("cleaning_transcript", "Cleaning transcript and removing repetition...", 72)
            with profiler.stage("chunk_summarization"):
                analysis_input = build_chunk_safe_analysis_input(transcript, query, progress_callback=progress_callback)
            chunks = chunk_text(analysis_input)
            try:
                with profiler.stage("topic_analysis_context"):
                    context = build_analysis_context(analysis_input, chunks, query)
            except Exception:
                context = " ".join(chunks[:3])
            try:
                with profiler.stage("analysis_generation"):
                    result = gemini_rag(context, query, source_url="")
                profiler.log_total()
                return result, transcript
            except Exception:
                profiler.log_total()
                return fallback_summary(transcript, query), transcript

    if not audio_source_path:
        raise RuntimeError("No audio source provided or audio extraction failed")

    if progress_callback:
        progress_callback("transcribing_audio", "Transcribing audio...", 35)
    try:
        with profiler.stage("transcription"):
            transcript = transcribe_audio_with_fallback(audio_source_path, cache_key=transcript_cache_key, progress_callback=progress_callback)
    except RuntimeError as exc:
        if "readable text" in str(exc).lower():
            profiler.log_total()
            return fallback_summary("", query), ""
        raise
    if not transcript.strip():
        profiler.log_total()
        return fallback_summary("", query), ""

    if progress_callback:
        progress_callback("cleaning_transcript", "Cleaning transcript and removing repetition...", 72)
    with profiler.stage("chunk_summarization"):
        analysis_input = build_chunk_safe_analysis_input(transcript, query, progress_callback=progress_callback)
    chunks = chunk_text(analysis_input)
    try:
        with profiler.stage("topic_analysis_context"):
            context = build_analysis_context(analysis_input, chunks, query)
    except Exception:
        context = " ".join(chunks[:3])

    try:
        with profiler.stage("analysis_generation"):
            result = gemini_rag(context, query, source_url="")
        profiler.log_total()
        return result, transcript
    except Exception:
        profiler.log_total()
        return fallback_summary(transcript, query), transcript


def analyze_text_content(transcript: str, query: str = "Summarize the content", *, progress_callback=None) -> tuple[dict[str, str], str]:
    profiler = StageProfiler("analyze-text")
    with profiler.stage("transcript_cleaning"):
        transcript = clean_source_text(transcript)
    if not transcript.strip():
        profiler.log_total()
        return fallback_summary("", query), ""

    if progress_callback:
        progress_callback("cleaning_transcript", "Cleaning transcript and source text...", 72)
    with profiler.stage("chunk_summarization"):
        analysis_input = build_chunk_safe_analysis_input(transcript, query, progress_callback=progress_callback)
    chunks = chunk_text(analysis_input)
    try:
        with profiler.stage("topic_analysis_context"):
            context = build_analysis_context(analysis_input, chunks, query)
    except Exception:
        context = " ".join(chunks[:3])

    try:
        with profiler.stage("analysis_generation"):
            result = gemini_rag(context, query, source_url="")
        profiler.log_total()
        return result, transcript
    except Exception:
        profiler.log_total()
        return fallback_summary(transcript, query), transcript


def analyze_youtube_via_transcription_fallback(url: str, query: str, *, progress_callback=None) -> tuple[dict[str, str], str]:
    with tempfile.TemporaryDirectory(prefix="media-analyzer-youtube-fallback-") as temp_dir:
        if progress_callback:
            progress_callback("audio_fallback", "Captions were not available. Transcribing audio...", 38)
        temp_audio_path = run_with_timeout(
            "Video download",
            min(120, YOUTUBE_FALLBACK_TIMEOUT),
            download_audio,
            url,
            temp_dir,
        )
        return run_with_timeout(
            "Transcription fallback",
            YOUTUBE_FALLBACK_TIMEOUT,
            analyze_media,
            None,
            temp_audio_path,
            query,
            transcript_cache_key=build_transcript_cache_key(url),
            progress_callback=progress_callback,
        )


def analyze_youtube_source(url: str, query: str = "Summarize the content", *, progress_callback=None) -> dict[str, object]:
    profiler = StageProfiler(f"youtube-analysis:{url}")
    analysis_cache_key = build_content_cache_key("youtube-analysis", f"{url}|{query}")
    with profiler.stage("analysis_cache_lookup"):
        cached = load_cached_analysis_result(analysis_cache_key)
    if cached:
        logger.info("youtube-analysis-cache-hit | %s", url)
        profiler.log_total()
        return cached

    metadata: dict[str, object] = {}
    transcript_text = ""
    source_context = ""
    direct_failure: Optional[Exception] = None

    if progress_callback:
        progress_callback("direct_gemini", "Analyzing video with Gemini...", 12)

    with profiler.stage("source_fetch_download"):
        metadata = run_with_timeout("YouTube metadata analysis", 35, fetch_youtube_metadata, url)

        try:
            with profiler.stage("analysis_generation"):
                logger.info("analysis-job | direct_gemini_start | url=%s", url)
                result = run_with_timeout(
                    "YouTube direct Gemini analysis",
                    YOUTUBE_DIRECT_ANALYSIS_TIMEOUT,
                    direct_gemini_youtube_analysis,
                    url,
                    query or "Summarize the content",
                    metadata,
                )
            source_context = build_youtube_source_text(metadata, "")
            if progress_callback:
                progress_callback("finalizing_output", "Finalizing headline, summary, and topic list...", 94)
            source_cache_key = build_content_cache_key("source-context", f"youtube|{url}|{source_context or result.get('summary', '')}")
            save_cached_source_content(
                source_cache_key,
                {"content": source_context or str(result.get("summary") or ""), "source_type": "youtube"},
            )
            payload = enrich_analysis(
                result,
                generate_article=False,
                source_context=source_context or str(result.get("summary") or ""),
                source_type="youtube",
                source_cache_key=source_cache_key,
            )
            payload["direct_analysis"] = True
            save_cached_analysis_result(analysis_cache_key, payload)
            profiler.log_total()
            return payload
        except Exception as direct_exc:
            direct_failure = direct_exc
            logger.warning(
                "youtube-direct-analysis-reason | url=%s | reason=%s",
                url,
                direct_gemini_failure_reason(direct_exc),
            )
            if progress_callback:
                progress_callback(
                    "captions_fallback",
                    "Direct video analysis took too long. Extracting transcript...",
                    28,
                )

        with tempfile.TemporaryDirectory(prefix="yt-direct-analysis-") as temp_dir:
            captions_error: Optional[Exception] = None
            try:
                with profiler.stage("transcription"):
                    transcript_text = run_with_timeout("YouTube transcript fetch", min(30, YOUTUBE_TRANSCRIPT_TIMEOUT), fetch_youtube_transcript_text, url)
            except Exception as transcript_exc:
                logger.warning(
                    "youtube-transcript-fetch-failed | url=%s | error_type=%s | error=%s",
                    url,
                    transcript_exc.__class__.__name__,
                    str(transcript_exc),
                )
                try:
                    with profiler.stage("transcription"):
                        transcript_text = run_with_timeout("YouTube subtitle fetch", min(30, YOUTUBE_SUBTITLE_TIMEOUT), fetch_youtube_subtitles_text, url, temp_dir)
                except Exception as subtitle_exc:
                    logger.warning(
                        "youtube-subtitle-fetch-failed | url=%s | error_type=%s | error=%s",
                        url,
                        subtitle_exc.__class__.__name__,
                        str(subtitle_exc),
                    )
                    captions_error = subtitle_exc

        if transcript_text.strip():
            if progress_callback:
                progress_callback("topic_generation", "Transcript found. Generating heading, summary, and topics...", 68)
            with profiler.stage("transcript_cleaning"):
                source_context = build_youtube_source_text(metadata, transcript_text)
            with profiler.stage("analysis_generation"):
                analysis_input = build_chunk_safe_analysis_input(source_context, query or "Summarize the content", progress_callback=progress_callback)
                chunks = chunk_text(analysis_input)
                try:
                    with profiler.stage("topic_analysis_context"):
                        context = build_analysis_context(analysis_input, chunks, query or "Summarize the content")
                except Exception:
                    context = " ".join(chunks[:3])
                result = gemini_rag(context, query or "Summarize the content", source_url=url)
            source_context = source_context
        else:
            if progress_callback:
                progress_callback("audio_fallback", "Captions were not available. Transcribing audio...", 38)
            if not remote_media_fallback_available():
                raise RuntimeError("Gemini direct analysis failed and no captions/transcript were available.") from captions_error or direct_exc
            result, cleaned_transcript = analyze_youtube_via_transcription_fallback(url, query or "Summarize the content", progress_callback=progress_callback)
            source_context = cleaned_transcript or source_context

    if progress_callback:
        progress_callback("finalizing_output", "Finalizing headline, summary, and topic list...", 94)
    source_cache_key = build_content_cache_key("source-context", f"youtube|{url}|{source_context or result.get('summary', '')}")
    save_cached_source_content(source_cache_key, {"content": source_context or str(result.get("summary") or ""), "source_type": "youtube"})
    payload = enrich_analysis(
        result,
        generate_article=False,
        source_context=source_context or str(result.get("summary") or ""),
        source_type="youtube",
        source_cache_key=source_cache_key,
    )
    payload["direct_analysis"] = direct_failure is None
    save_cached_analysis_result(analysis_cache_key, payload)
    profiler.log_total()
    return payload


def analyze_url_source(url: str, query: str = "Summarize the content", *, progress_callback=None) -> dict[str, object]:
    profiler = StageProfiler(f"url-analysis:{url}")
    transcript_override: Optional[str] = None
    temp_audio_path: Optional[str] = None
    source_type = source_kind_from_url(url)

    if source_type == "unsupported":
        raise RuntimeError("Direct image links are not supported yet. Please use a webpage, YouTube/video URL, or upload an audio/video file.")

    if source_type == "youtube":
        payload = analyze_youtube_source(url, query or "Summarize the content", progress_callback=progress_callback)
        profiler.log_total()
        return payload

    if source_type == "web-url":
        if progress_callback:
            progress_callback("downloading_source", "Fetching website content...", 15)
        with profiler.stage("source_fetch_download"):
            webpage = extract_webpage_content(url)
        if progress_callback:
            progress_callback("analyzing_topic", "Analyzing topic and key ideas...", 78)
        result, cleaned_transcript = analyze_text_content(
            str(webpage.get("content") or ""),
            query=query or "Summarize the content",
            progress_callback=progress_callback,
        )
        source_cache_key = build_content_cache_key("source-context", f"web-url|{url}|{cleaned_transcript}")
        save_cached_source_content(source_cache_key, {"content": cleaned_transcript, "source_type": "web-url"})
        payload = enrich_analysis(
            result,
            generate_article=False,
            source_context=cleaned_transcript or str(webpage.get("content") or ""),
            source_type="web-url",
            source_cache_key=source_cache_key,
        )
        profiler.log_total()
        return payload

    with tempfile.TemporaryDirectory(prefix="media-analyzer-url-") as temp_dir:
        if source_type == "youtube":
            if progress_callback:
                progress_callback("downloading_source", "Downloading source audio from YouTube...", 12)
            try:
                with profiler.stage("source_fetch_download"):
                    temp_audio_path = download_audio(url, temp_dir)
            except Exception as download_exc:
                raise RuntimeError(f"Video could not be downloaded. {download_exc}") from download_exc
        else:
            if progress_callback:
                progress_callback("downloading_source", "Downloading source audio...", 12)
            if not remote_media_fallback_available():
                raise RuntimeError("Direct remote media downloads are disabled in fast mode. Please upload the audio/video file instead.")
            with profiler.stage("source_fetch_download"):
                temp_audio_path = download_audio(url, temp_dir)

        if transcript_override is not None:
            result, cleaned_transcript = analyze_text_content(
                clean_source_text(transcript_override),
                query=query or "Summarize the content",
                progress_callback=progress_callback,
            )
        else:
            result, cleaned_transcript = analyze_media(
                local_audio_path=temp_audio_path,
                query=query or "Summarize the content",
                transcript_cache_key=build_transcript_cache_key(url),
                progress_callback=progress_callback,
            )

    if progress_callback:
        progress_callback("finalizing_output", "Finalizing output and preparing topics...", 96)
    source_cache_key = build_content_cache_key("source-context", f"{source_type}|{url}|{cleaned_transcript or transcript_override or result.get('summary', '')}")
    save_cached_source_content(source_cache_key, {"content": cleaned_transcript or transcript_override or result.get("summary", ""), "source_type": source_type})
    payload = enrich_analysis(
        result,
        generate_article=False,
        source_context=cleaned_transcript or transcript_override or result.get("summary", ""),
        source_type="youtube" if source_type == "youtube" else "media-url",
        source_cache_key=source_cache_key,
    )
    profiler.log_total()
    return payload


def serialize_job_status(job: Job) -> dict[str, object]:
    progress_payload = {
        "stage": str(job.meta.get("stage") or "queued"),
        "message": str(job.meta.get("message") or "We are processing your source in the background."),
        "progress": int(job.meta.get("progress") or 0),
    }
    status = job.get_status(refresh=True)
    now = time.time()
    raw_created_at_ts = job.meta.get("created_at_ts")
    raw_updated_at_ts = job.meta.get("updated_at_ts")
    created_at_ts = float(raw_created_at_ts) if raw_created_at_ts is not None else now
    updated_at_ts = float(raw_updated_at_ts) if raw_updated_at_ts is not None else created_at_ts
    age_seconds = max(0, int(now - created_at_ts))
    stale_seconds = max(0, int(now - updated_at_ts))

    if status not in {"finished", "failed"}:
        if age_seconds >= ANALYSIS_JOB_MAX_SECONDS:
            logger.warning(
                "analysis-job-timeout | job_id=%s | status=%s | age=%ss | stale=%ss | stage=%s",
                job.id,
                status,
                age_seconds,
                stale_seconds,
                progress_payload["stage"],
            )
            return {
                "success": False,
                "status": "failed",
                "reason": "timeout",
                "message": "We could not extract reliable content from this video. Please upload the video/audio file or try another link.",
                "error": "We could not extract reliable content from this video. Please upload the video/audio file or try another link.",
                "progress": {
                    **progress_payload,
                    "stage": "failed",
                    "message": "We could not extract reliable content from this video. Please upload the video/audio file or try another link.",
                },
            }
        if progress_payload["stage"] == "direct_gemini" and stale_seconds >= YOUTUBE_DIRECT_ANALYSIS_TIMEOUT:
            logger.info(
                "analysis-job-stage-promote | job_id=%s | from=%s | to=%s | stale=%ss",
                job.id,
                progress_payload["stage"],
                "captions_fallback",
                stale_seconds,
            )
            progress_payload = {
                "stage": "captions_fallback",
                "message": "Direct video analysis took too long. Extracting transcript...",
                "progress": max(progress_payload["progress"], 28),
            }
        elif progress_payload["stage"] == "captions_fallback" and stale_seconds >= max(30, YOUTUBE_TRANSCRIPT_TIMEOUT):
            logger.info(
                "analysis-job-stage-promote | job_id=%s | from=%s | to=%s | stale=%ss",
                job.id,
                progress_payload["stage"],
                "audio_fallback",
                stale_seconds,
            )
            progress_payload = {
                "stage": "audio_fallback",
                "message": "Captions were not available. Transcribing audio...",
                "progress": max(progress_payload["progress"], 38),
            }
    if status == "finished":
        return {
            "success": True,
            "status": "completed",
            "result": job.result,
            "progress": {**progress_payload, "progress": 100, "stage": "completed", "message": "Analysis completed."},
        }

    if status == "failed":
        last_line = ""
        if job.exc_info:
            last_line = job.exc_info.strip().splitlines()[-1]
        return {
            "success": False,
            "status": "failed",
            "reason": "error",
            "error": last_line or "Background job failed.",
            "progress": progress_payload,
        }

    return {
        "success": True,
        "status": status,
        "queued": True,
        "jobId": job.id,
        "progress": progress_payload,
    }


def build_articles_response(payload: ArticlesRequest) -> dict[str, object]:
    profiler = StageProfiler("generate-articles-endpoint")
    selected_topics = [topic.strip() for topic in payload.selected_topics if topic.strip()]
    if not selected_topics:
        raise HTTPException(status_code=400, detail="Select at least one topic before generating articles.")

    source_context = payload.source_context
    if payload.source_cache_key:
        with profiler.stage("source_cache_lookup"):
            cached_source = load_cached_source_content(payload.source_cache_key)
        if isinstance(cached_source, dict) and isinstance(cached_source.get("content"), str):
            source_context = str(cached_source["content"])

    selected_topic_summary = ""
    selected_topic_points: list[dict[str, str]] = []
    if isinstance(payload.selected_topic_details, dict):
        selected_topic_summary = str(
            payload.selected_topic_details.get("summary")
            or payload.selected_topic_details.get("explanation")
            or ""
        ).strip()
        raw_points = payload.selected_topic_details.get("points", [])
        if isinstance(raw_points, list):
            selected_topic_points = [
                {
                    "label": str(item.get("label", "")).strip(),
                    "description": str(item.get("description", "")).strip(),
                }
                for item in raw_points
                if isinstance(item, dict)
            ]

    article_cache_key = build_content_cache_key(
        "article-result",
        json.dumps(
            {
                "headline": payload.headline.strip() or "Media summary",
                "summary": payload.summary.strip(),
                "selected_topics": selected_topics,
                "article_count": payload.article_count,
                "article_type": payload.article_type,
                "target_audience": payload.target_audience,
                "source_cache_key": payload.source_cache_key,
                "source_context": source_context[:2500],
                "selected_topic_summary": selected_topic_summary,
                "selected_topic_points": selected_topic_points,
            },
            sort_keys=True,
        ),
    )
    with profiler.stage("article_cache_lookup"):
        cached_articles = load_cached_article_result(article_cache_key)
    if isinstance(cached_articles, dict):
        profiler.log_total()
        return cached_articles

    base_result = {
        "headline": payload.headline.strip() or "Media summary",
        "summary": payload.summary.strip(),
    }
    with profiler.stage("article_generation"):
        enriched_result = enrich_analysis(
            base_result,
            generate_article=True,
            selected_topics=selected_topics,
            article_count=payload.article_count,
            article_type=payload.article_type,
            target_audience=payload.target_audience,
            source_context="\n".join(part for part in [source_context, selected_topic_summary, "\n".join(f"{point['label']}: {point['description']}" for point in selected_topic_points if point.get('label') or point.get('description'))] if part).strip(),
            source_cache_key=payload.source_cache_key,
        )
    if payload.topics:
        enriched_result["topics"] = payload.topics
    save_cached_article_result(article_cache_key, enriched_result)
    profiler.log_total()
    return enriched_result


def render_article_to_docx_bytes(title: str, topic: str, content_html: str) -> bytes:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("DOCX export is not available on this deployment yet.") from exc

    document = Document()
    document.add_heading(title, level=1)
    if topic:
        document.add_paragraph(f"Topic: {topic}")
    for line in re.split(r"\n+", postprocess_article_html(content_html)):
        line = line.strip()
        if not line:
            continue
        if line.startswith("<h2>"):
            document.add_heading(strip_html_tags(line), level=1)
        elif line.startswith("<h3>"):
            document.add_heading(strip_html_tags(line), level=2)
        elif line.startswith("<li>"):
            document.add_paragraph(strip_html_tags(line), style="List Bullet")
        else:
            document.add_paragraph(strip_html_tags(line))
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_article_to_pdf_bytes(title: str, topic: str, content_html: str) -> bytes:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except Exception as exc:
        raise RuntimeError("PDF export is not available on this deployment yet.") from exc

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = [Paragraph(title, styles["Title"])]
    if topic:
        story.append(Paragraph(f"<b>Topic:</b> {topic}", styles["BodyText"]))
        story.append(Spacer(1, 8))
    for line in re.split(r"\n+", postprocess_article_html(content_html)):
        stripped = line.strip()
        if not stripped:
            continue
        text = strip_html_tags(stripped)
        if stripped.startswith("<h2>"):
            story.append(Paragraph(text, styles["Heading1"]))
        elif stripped.startswith("<h3>"):
            story.append(Paragraph(text, styles["Heading2"]))
        else:
            story.append(Paragraph(text, styles["BodyText"]))
        story.append(Spacer(1, 6))
    doc.build(story)
    return buffer.getvalue()


async def save_uploaded_file(upload: UploadFile, destination: Path) -> None:
    content = await upload.read()
    destination.write_bytes(content)


@app.get("/")
def index():
    index_file = FRONTEND_DIR / "index.html"
    if index_file.exists():
        return FileResponse(index_file)
    return HTMLResponse("<h1>ZENPEN is running</h1><p>Frontend assets are not available in this deployment bundle.</p>")


@app.get("/api/health")
def health() -> JSONResponse:
    return JSONResponse({"status": "ok", "dependencies": dependency_status()})


@app.get("/api/config")
def config() -> JSONResponse:
    return JSONResponse(get_api_config())


@app.post("/api/auth/signup")
def signup(payload: SignupRequest) -> JSONResponse:
    if not supabase_is_configured():
        raise HTTPException(status_code=503, detail="Supabase authentication is not configured.")
    if len(payload.password) < 6:
        raise HTTPException(status_code=400, detail="Password must be at least 6 characters long.")

    create_supabase_user(payload.name.strip(), payload.email.strip().lower(), payload.password)
    session = sign_in_supabase_user(payload.email.strip().lower(), payload.password)
    return JSONResponse({"success": True, "session": session})


@app.post("/api/auth/login")
def login(payload: LoginRequest) -> JSONResponse:
    if not supabase_is_configured():
        raise HTTPException(status_code=503, detail="Supabase authentication is not configured.")

    session = sign_in_supabase_user(payload.email.strip().lower(), payload.password)
    return JSONResponse({"success": True, "session": session})


@app.get("/api/history")
def history(request: Request) -> JSONResponse:
    user = resolve_supabase_user(request)
    if user:
        try:
            return JSONResponse(load_history_from_supabase(str(user["id"])))
        except Exception:
            pass
    return JSONResponse(load_history())


@app.get("/api/jobs/{job_id}")
def job_status(job_id: str) -> JSONResponse:
    job = fetch_job(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Analysis job not found.")
    payload = serialize_job_status(job)
    return JSONResponse(payload, status_code=200)


@app.post("/api/analyze-youtube")
def analyze_youtube_endpoint(payload: AnalyzeYouTubeRequest, request: Request) -> JSONResponse:
    url = payload.url.strip()
    if not url:
        raise HTTPException(status_code=400, detail="Please provide a YouTube URL.")
    if not is_youtube_url(url):
        raise HTTPException(status_code=400, detail="Only YouTube links are supported by this endpoint.")

    try:
        user = resolve_supabase_user(request)
        if background_url_jobs_available():
            job = enqueue_url_analysis(
                url=url,
                query=payload.query or "Give breaking news and main points",
                user_id=str(user["id"]) if user else None,
            )
            return JSONResponse(
                {
                    "success": True,
                    "queued": True,
                    "jobId": job.id,
                    "status": "queued",
                    "progress": {
                        "stage": "direct_gemini",
                        "message": "Analyzing video with Gemini...",
                        "progress": 6,
                    },
                },
                status_code=202,
            )

        result = analyze_youtube_source(url, payload.query or "Give breaking news and main points")
        return JSONResponse({"success": True, "result": result})
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("youtube-analyze-failed")
        return JSONResponse({"success": False, "error": map_public_error_message(str(exc))}, status_code=500)


@app.post("/api/analyze-source")
async def analyze_source_alias_endpoint(
    request: Request,
    url: Optional[str] = Form(None),
    query: Optional[str] = Form("Summarize the content"),
    file: Optional[UploadFile] = File(None),
    generate_article: bool = Form(False),
    article_count: int = Form(1),
    selected_topics: Optional[str] = Form(None),
    article_type: str = Form("Blog Article"),
    target_audience: str = Form("General readers"),
) -> JSONResponse:
    return await analyze_endpoint(
        request=request,
        url=url,
        query=query,
        file=file,
        generate_article=generate_article,
        article_count=article_count,
        selected_topics=selected_topics,
        article_type=article_type,
        target_audience=target_audience,
    )


@app.post("/api/articles")
def generate_articles_endpoint(payload: ArticlesRequest) -> JSONResponse:
    enriched_result = build_articles_response(payload)
    return JSONResponse({"success": True, "result": enriched_result})


@app.post("/api/generate-article")
def generate_article_alias_endpoint(payload: ArticlesRequest) -> JSONResponse:
    enriched_result = build_articles_response(payload)
    return JSONResponse({"success": True, "result": enriched_result})


@app.post("/api/articles/publish")
def publish_articles_endpoint(request: Request, payload: PublishArticleRequest) -> JSONResponse:
    user = resolve_supabase_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Please log in before publishing a draft.")

    result_payload = {
        "headline": payload.headline,
        "summary": payload.summary,
        "topics": payload.topics,
        "articles": payload.articles,
    }
    persist_analysis_to_supabase(
        user_id=str(user["id"]),
        source_type=payload.source_type,
        source_url=payload.source_url,
        source_file_name=payload.source_file_name,
        source_mime_type=payload.source_mime_type,
        query=payload.query,
        result=result_payload,
        selected_topics=payload.selected_topics,
    )
    return JSONResponse({"success": True, "message": "Draft saved successfully."})


@app.post("/api/articles/export")
def export_article_endpoint(payload: ExportArticleRequest):
    profiler = StageProfiler(f"export:{payload.format.lower().strip()}")
    title = payload.title.strip() or "Generated Article"
    topic = payload.topic.strip()
    content_html = payload.content_html or ""
    export_format = payload.format.lower().strip()

    if export_format == "txt":
        with profiler.stage("export_preparation"):
            text = f"{title}\n\nTopic: {topic}\n\n{strip_html_tags(content_html)}"
        profiler.log_total()
        return StreamingResponse(
            BytesIO(text.encode("utf-8")),
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{build_article_slug(title)}.txt"'},
        )
    if export_format == "html":
        with profiler.stage("export_preparation"):
            html = postprocess_article_html(content_html)
        profiler.log_total()
        response = HTMLResponse(html)
        response.headers["Content-Disposition"] = f'attachment; filename="{build_article_slug(title)}.html"'
        return response
    if export_format == "docx":
        with profiler.stage("export_preparation"):
            data = render_article_to_docx_bytes(title, topic, content_html)
        profiler.log_total()
        return StreamingResponse(
            BytesIO(data),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{build_article_slug(title)}.docx"'},
        )
    if export_format == "pdf":
        with profiler.stage("export_preparation"):
            data = render_article_to_pdf_bytes(title, topic, content_html)
        profiler.log_total()
        return StreamingResponse(
            BytesIO(data),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{build_article_slug(title)}.pdf"'},
        )
    raise HTTPException(status_code=400, detail="Unsupported export format.")


@app.post("/api/analyze")
async def analyze_endpoint(
    request: Request,
    url: Optional[str] = Form(None),
    query: Optional[str] = Form("Summarize the content"),
    file: Optional[UploadFile] = File(None),
    generate_article: bool = Form(False),
    article_count: int = Form(1),
    selected_topics: Optional[str] = Form(None),
    article_type: str = Form("Blog Article"),
    target_audience: str = Form("General readers"),
) -> JSONResponse:
    profiler = StageProfiler("analyze-endpoint")
    if not file and not url:
        raise HTTPException(status_code=400, detail="Please provide a URL or upload an audio/video file")
    if file:
        content_type = (file.content_type or "").lower()
        suffix = Path(file.filename or "").suffix.lower()
        mime_supported = any(content_type.startswith(prefix) for prefix in SUPPORTED_UPLOAD_MIME_PREFIXES)
        suffix_supported = suffix in SUPPORTED_MEDIA_EXTENSIONS
        if not (mime_supported or suffix_supported):
            raise HTTPException(status_code=400, detail="Unsupported file type. Please upload an audio or video file.")

    try:
        user = resolve_supabase_user(request)
        if url and not file and background_url_jobs_available():
            job = enqueue_url_analysis(
                url=url,
                query=query or "Summarize the content",
                user_id=str(user["id"]) if user else None,
            )
            is_youtube = is_youtube_url(url)
            return JSONResponse(
                {
                    "success": True,
                    "queued": True,
                    "jobId": job.id,
                    "status": "queued",
                    "progress": {
                        "stage": "direct_gemini" if is_youtube else "downloading_source",
                        "message": "Analyzing video with Gemini..." if is_youtube else "Preparing source...",
                        "progress": 5,
                    },
                },
                status_code=202,
            )
        with tempfile.TemporaryDirectory(prefix="media-analyzer-") as temp_dir:
            temp_dir_path = Path(temp_dir)
            temp_audio_path: Optional[str] = None
            topic_list = [topic.strip() for topic in (selected_topics or "").split(",") if topic.strip()]
            source_context = ""
            source_cache_key = ""
            detected_source_type = "upload" if file else source_kind_from_url(url or "")

            if file:
                suffix = Path(file.filename or "upload.wav").suffix.lower() or ".wav"
                upload_path = temp_dir_path / f"upload{suffix}"
                with profiler.stage("source_fetch_download"):
                    await save_uploaded_file(file, upload_path)
                if suffix in {".mp4", ".mov", ".mkv", ".avi", ".webm"}:
                    with profiler.stage("source_extract_audio"):
                        temp_audio_path = extract_audio_from_video(
                            str(upload_path),
                            str(temp_dir_path / "extracted_audio.wav"),
                        )
                else:
                    temp_audio_path = str(upload_path)
                result, cleaned_transcript = analyze_media(
                    local_audio_path=temp_audio_path,
                    query=query or "Summarize the content",
                )
                source_context = cleaned_transcript or str(result.get("summary") or "")
                source_cache_key = build_content_cache_key("source-context", f"upload|{file.filename}|{source_context}")
                save_cached_source_content(source_cache_key, {"content": source_context, "source_type": "upload"})
                enriched_result = enrich_analysis(
                    result,
                    generate_article=generate_article,
                    selected_topics=topic_list or None,
                    article_count=article_count,
                    article_type=article_type,
                    target_audience=target_audience,
                    source_context=source_context,
                    source_type="upload",
                    source_cache_key=source_cache_key,
                )
            elif url:
                enriched_result = analyze_url_source(url, query or "Summarize the content")
                source_context = str(enriched_result.get("source_context_preview") or enriched_result.get("summary") or "")
                source_cache_key = str(enriched_result.get("source_cache_key") or "")
                if generate_article:
                    existing_topics = list(enriched_result.get("topics") or [])
                    base_result = {
                        "headline": str(enriched_result.get("headline") or "Media summary"),
                        "summary": str(enriched_result.get("summary") or ""),
                    }
                    enriched_result = enrich_analysis(
                        base_result,
                        generate_article=True,
                        selected_topics=topic_list or None,
                        article_count=article_count,
                        article_type=article_type,
                        target_audience=target_audience,
                        source_context=source_context,
                        source_type=detected_source_type,
                        source_cache_key=source_cache_key,
                    )
                    enriched_result["topics"] = existing_topics
            source = "uploaded-file" if file else ("youtube-url" if detected_source_type == "youtube" else "remote-url")
            source_type = "upload" if file else "url"
            result_for_history = {
                "headline": str(enriched_result.get("headline") or "Media summary"),
                "summary": str(enriched_result.get("summary") or ""),
            }
            add_history_entry(result_for_history, source)
            if user:
                try:
                    persist_analysis_to_supabase(
                        user_id=str(user["id"]),
                        source_type=source_type,
                        source_url=url,
                        source_file_name=file.filename if file else None,
                        source_mime_type=file.content_type if file else None,
                        query=query or "Summarize the content",
                        result=enriched_result,
                        selected_topics=topic_list,
                    )
                except Exception:
                    pass
            profiler.log_total()
            return JSONResponse({"success": True, "result": enriched_result})
    except HTTPException:
        raise
    except Exception as exc:
        print(f"[analyze_endpoint] {exc}", file=sys.stderr)
        return JSONResponse({"success": False, "error": map_public_error_message(str(exc))}, status_code=500)


@app.get("/{full_path:path}")
def frontend_routes(full_path: str):
    if full_path.startswith(("api/", "static/")):
        raise HTTPException(status_code=404, detail="Not found")
    index_file = FRONTEND_DIR / "index.html"
    if index_file.exists():
        return FileResponse(index_file)
    return HTMLResponse("<h1>ZENPEN is running</h1><p>Frontend assets are not available in this deployment bundle.</p>")
